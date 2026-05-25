import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_weekly_tactical_docx(intel_data):
    doc = Document()
    
    # --- Title (Centred) ---
    title = doc.add_heading(intel_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Date & Classification (Centred) ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"DATE: {intel_data['date']} | {intel_data['classification']}")
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0, 112, 192) 
    run.bold = True
    
    doc.add_paragraph() # Spacer
    
    # --- BLUF (Bottom Line Up Front) ---
    bluf_head = doc.add_heading('BLUF (Bottom Line Up Front)', level=1)
    bluf_para = doc.add_paragraph()
    bluf_run = bluf_para.add_run(intel_data['bluf'])
    bluf_run.bold = True
    bluf_run.font.color.rgb = RGBColor(192, 0, 0) # Dark red for BLUF
    
    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(intel_data['executive_summary'])
    
    # --- Threat Narrative ---
    doc.add_heading('Threat Narrative', level=1)
    doc.add_paragraph(intel_data['threat_narrative'])
    
    # --- Risk Assessment ---
    doc.add_heading('Risk Assessment', level=1)
    doc.add_paragraph(intel_data['risk_assessment'])
    
    # --- Tactical Indicators (Bullets) ---
    doc.add_heading('Tactical Indicators', level=1)
    for indicator in intel_data['tactical_indicators']:
        clean_indicator = indicator.replace('**', '') 
        doc.add_paragraph(clean_indicator, style='List Bullet')
        
    doc.add_paragraph() # Spacer
    
    # --- Predictive Analysis ---
    doc.add_heading('Predictive Analysis', level=1)
    doc.add_paragraph(intel_data['predictive_analysis'])
    
    # --- Strategic Recommendations ---
    doc.add_heading('Strategic Recommendations', level=1)
    doc.add_paragraph(intel_data['recommendations'])
    
    # --- Save to Memory Buffer ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer