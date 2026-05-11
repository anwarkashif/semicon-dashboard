import re
from io import BytesIO
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import streamlit as st

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000EE')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None: tblPr.remove(tblBorders)
    new_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        new_borders.append(border)
    tblPr.append(new_borders)

def add_word_data_table(doc, heading_text, data_list):
    if not data_list: return
    safe_list = []
    for item in data_list:
        if isinstance(item, dict):
            safe_list.append(item)
        else:
            safe_list.append({"Extracted Information": str(item)})
            
    if not safe_list: return
    if len(safe_list) == 1 and "No " in str(list(safe_list[0].values())[0]): return
    
    p_heading = doc.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_heading.paragraph_format.line_spacing = 1.5
    p_heading.paragraph_format.space_before = Pt(6) 
    p_heading.paragraph_format.space_after = Pt(0) 
    r_head = p_heading.add_run(heading_text)
    r_head.bold = True
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(12)

    headers = list(safe_list[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER 
    table.style = 'Table Grid'
    set_table_borders(table) 
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r.bold = True
                
    for item in safe_list:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            row_cells[i].text = str(item.get(h, ''))
            for p in row_cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
    doc.add_paragraph() 

@st.cache_data
def create_landscape_word(text_sections, final_text, actions_data, brief_date, fund_data, market_data, risk_data, sources_data, text_ews):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    header = section.header
    header.paragraphs[0].text = "" 
    r_head = header.paragraphs[0].add_run("Kashif Anwar, SemicoN Dashboard Brief")
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(11)
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer = section.footer
    footer.paragraphs[0].text = ""
    r_foot = footer.paragraphs[0].add_run("Kashif Anwar, SemicoN Dashboard Brief")
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(11)
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mast_table = doc.add_table(rows=1, cols=2)
    mast_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_logo = mast_table.cell(0, 0)
    cell_logo.width = Inches(2.0)
    cell_text = mast_table.cell(0, 1)
    cell_text.width = Inches(6.0)
    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img = Image.open("logo.jpg").convert("RGBA")
        min_dim = min(img.size)
        img = img.crop((0, 0, min_dim, min_dim))
        mask = Image.new('L', img.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, min_dim, min_dim), fill=255)
        result = Image.new('RGBA', img.size, (255, 255, 255, 0))
        buf = BytesIO()
        result.save(buf, format="PNG")
        buf.seek(0)
        p_logo.add_run().add_picture(buf, width=Inches(1.5))
    except Exception: pass 
        
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_text = p_text.add_run("SemicoN Dashboard – A Semicon News Dashboard")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(20)
    run_text.bold = True
    run_text.font.color.rgb = RGBColor(255, 0, 0) 
    doc.add_paragraph() 
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.5
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12) 
    run_title = p_title.add_run(f'SemicoN Weekly Brief - {brief_date}')
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.name = 'Times New Roman'
    
    p_author = doc.add_paragraph()
    p_author.paragraph_format.line_spacing = 1.5
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after = Pt(12)
    run_author = p_author.add_run("Prepared By: Kashif Anwar")
    run_author.bold = True
    run_author.font.size = Pt(12)
    run_author.font.name = 'Times New Roman'
    
    section_titles = [
        "Executive Summary",
        "Global Foundry Market & Geopolitical Positioning",
        "AI Chip Demand, Manufacturing & Processing",
        "Critical Minerals: Rare Earth Reserves & Supply Chains",
        "Export Controls & Geopolitical Impact",
        "AI, Chips and Rare Earth in Military and Outer Space Domain",
        "Lithography Chokepoints & State Actions",
        "India: Domestic & Strategic Developments",
        "West Asia/Middle East: Domestic & Strategic Developments"
    ]

    for i, text_data in enumerate(text_sections):
        
        if i == 0 and text_ews and text_ews.strip() != "":
            p_ews_head = doc.add_paragraph()
            p_ews_head.paragraph_format.space_before = Pt(12)
            r_ews_head = p_ews_head.add_run("🚨 Early Warning & Red Flags")
            r_ews_head.bold = True
            r_ews_head.font.color.rgb = RGBColor(255, 0, 0)
            r_ews_head.font.name = 'Times New Roman'
            r_ews_head.font.size = Pt(14)
            
            p_ews = doc.add_paragraph()
            r_ews = p_ews.add_run(text_ews)
            r_ews.font.name = 'Times New Roman'
            r_ews.font.size = Pt(12)
            
        if not text_data or text_data.strip() == "": continue

        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(18)
        r_head = p_head.add_run(section_titles[i])
        r_head.bold = True
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.color.rgb = RGBColor(0, 102, 204)

        for line in text_data.split('\n'):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            
            if line.startswith('**') and line.endswith('**'):
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.replace('**', ''))
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif line.startswith('* '):
                p.style = 'List Bullet'
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line[2:].replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line.replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if i == 1: add_word_data_table(doc, "Strategic Investments & Funding", fund_data)
        if i == 3: add_word_data_table(doc, "Market & Geopolitical Impact", market_data)
        if i == 4: add_word_data_table(doc, "Supply Chain Risk Analysis", risk_data)
                
    add_word_data_table(doc, 'Recent State Actions', actions_data)

    if final_text and final_text.strip() != "":
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(18)
        r_head = p_head.add_run("Strategic Conclusion")
        r_head.bold = True
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.color.rgb = RGBColor(0, 102, 204)

        for line in final_text.split('\n'):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            
            if line.startswith('**') and line.endswith('**'):
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.replace('**', ''))
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif line.startswith('* '):
                p.style = 'List Bullet'
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line[2:].replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line.replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
    if sources_data:
        doc.add_paragraph() 
        p_source_head = doc.add_paragraph()
        p_source_head.paragraph_format.line_spacing = 1.5
        p_source_head.paragraph_format.space_before = Pt(12)
        r_source = p_source_head.add_run("Verified Intelligence Sources")
        r_source.bold = True
        r_source.font.name = 'Times New Roman'
        r_source.font.size = Pt(14)
        
        for src in sources_data:
            p_src = doc.add_paragraph()
            p_src.style = 'List Bullet'
            p_src.paragraph_format.space_after = Pt(6)
            p_src.paragraph_format.line_spacing = 1.0
            add_hyperlink(p_src, src['url'], src['title'])
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer