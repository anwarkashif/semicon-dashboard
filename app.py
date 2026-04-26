import streamlit as st
import pandas as pd
import os
import glob
import json
import re
import time
import sys
import math
import streamlit.components.v1 as components
from io import BytesIO
from PIL import Image, ImageDraw
from datetime import datetime, timezone
import plotly.express as px
import plotly.graph_objects as go
import pydeck as pdk
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import base64

# ==========================================
# 1. PAGE CONFIGURATION (MUST BE AT VERY TOP)
# ==========================================
st.set_page_config(
    page_title="SemicoN Dashboard", 
    page_icon="🛡️",  # <-- FIX 1: Use emoji to prevent fleeting startup errors
    layout="wide", 
    initial_sidebar_state="collapsed"
)

# FIX 1 & 3: Initialize session state immediately so routing doesn't flash
if 'role' not in st.session_state: 
    st.session_state['role'] = None

# ==========================================
# EARLY CSS ROUTING (KILLS THE FLASH/BLINK BUG)
# ==========================================
if st.session_state['role'] is None:
    st.markdown("""
    <style>
        /* SAFE UI HIDING */
        [data-testid="stToolbar"], [data-testid="stHeader"], [data-testid="collapsedControl"], [data-testid="stSidebar"] { 
            display: none !important; 
        }
        footer { visibility: hidden; }

        /* FIX 1: Use 100dvh for virtual keyboards and allow safe overflow */
        html, body, .stApp { min-height: 100dvh; margin: 0; background-color: #000000; overflow-x: hidden; }

        /* FIX 1: Prevent Mobile Browser Auto-Zoom on inputs */
        input[type="text"], input[type="password"] {
            font-size: 16px !important; 
        }

        /* --- 1. THE PURE CSS LEFT PANEL --- */
        .fixed-left-panel {
            position: fixed; top: 0; left: 0; width: 50vw; height: 100vh;
            background: linear-gradient(135deg, #0f172a, #1e293b, #020617);
            padding: 60px; display: flex; flex-direction: column;
            justify-content: center; z-index: 100; box-sizing: border-box;
            border-right: 1px solid #222;
        }
        .fixed-left-panel h1 { font-size: 3.2rem; font-weight: 300; line-height: 1.2; margin-bottom: 10px; color: white;}
        .fixed-left-panel span { font-weight: 700; color: #facc15; }
        .fixed-left-panel p { color: #94a3b8; font-size: 1.2rem; margin-top: 10px; }

        /* --- 2. THE RIGHT PANEL (STREAMLIT NATIVE CONTAINER) --- */
        .block-container {
            margin-left: 50vw !important; width: 50vw !important; max-width: 50vw !important;
            height: 100vh !important; padding: 0 15% !important;
            display: flex !important; flex-direction: column !important;
            justify-content: center !important;
        }

        /* --- 3. LOGIN HEADER --- */
        .login-header { display: flex; flex-direction: column; align-items: center; justify-content: center; margin-bottom: 25px; text-align: center; }
        .login-logo { width: 300px; max-width: 100%; margin-bottom: 5px; image-rendering: crisp-edges; }
        .login-header h2 { margin: 0; font-size: 32px; color: white; }
        .login-header p { margin-top: 5px; color: #aaa; font-size: 14px; }

        .stButton>button[kind="secondary"] { width: 100%; font-weight: bold; height: 45px; }

        /* --- LOGIN BUTTON YELLOW OVERRIDE --- */
        button[kind="primaryFormSubmit"],
        button[kind="primary"],
        div[data-testid="stFormSubmitButton"] button {
            background-color: #facc15 !important;
            color: #000000 !important;
            border: none !important;
            width: 100% !important;
            height: 45px !important;
            font-weight: bold !important;
            border-radius: 8px !important;
        }
        button[kind="primaryFormSubmit"]:hover,
        button[kind="primary"]:hover,
        div[data-testid="stFormSubmitButton"] button:hover {
            background-color: #eab308 !important;
        }

        /* --- 5. MOBILE OVERRIDES --- */
        @media screen and (max-width: 900px) {
            html, body, .stApp { overflow: auto; }
            .fixed-left-panel { display: none !important; }
            .block-container {
                margin-left: 0 !important; width: 100vw !important; max-width: 100vw !important;
                padding: 2rem !important; position: relative !important;
                height: auto !important; min-height: 100vh !important;
            }
        }
    </style>
    """, unsafe_allow_html=True)
else:
    # Pre-load the normal dashboard layout before DOM renders
    st.markdown("""
    <style>
        /* AGGRESSIVELY HIDE ALL LOGIN ELEMENTS TO PREVENT 1-SECOND GHOSTING FLASH */
        .fixed-left-panel, .login-header { 
            display: none !important; 
            visibility: hidden !important; 
            opacity: 0 !important; 
            z-index: -999 !important; 
            height: 0px !important;
        } 
        
        .block-container {
            margin-left: auto !important; margin-right: auto !important;
            /* FIX 2: Changed from 100vw to 100% so it respects the sidebar width */
            width: 100% !important; max-width: 100% !important;
            display: block !important; /* CRITICAL FIX: Releases the flexbox lock from the login screen */
            height: auto !important;   /* CRITICAL FIX: Allows dashboard to scroll normally */
        }
        [data-testid="stToolbar"], [data-testid="collapsedControl"] { display: flex !important; }
    </style>
    """, unsafe_allow_html=True)

# --- App Configuration ---
MAINTENANCE_MODE = False
if MAINTENANCE_MODE:
    st.markdown("""<style>, [data-testid="stSidebar"] { display: none; }</style>""", unsafe_allow_html=True)
    st.markdown("<div style='margin-top: 15vh;'></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        try:
            st.image("logo.jpg", width=300)
        except:
            pass
        st.warning("⚠️ **Warning: Work in Progress.** Please wait, the Dashboard will be live soon.")
    st.stop()

# --- App Initialization & Folder Setup ---
os.makedirs('data', exist_ok=True)
os.makedirs('trash', exist_ok=True)

# ==========================================
# Consider moving this to an environment variable in the future: os.environ.get("MAPBOX_TOKEN")
MAPBOX_PUBLIC_TOKEN = "pk.eyJ1Ijoia2FzaGlmYW53YXIiLCJhIjoiY21td2loemd2Mm10MzJycXh4aTd1YjZtdCJ9.EN4o_kXPmA8ScOimJyf53A"
# ==========================================

# --- API Setup ---
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None
model_name = 'gemini-2.5-flash'

# --- Comprehensive Global ISO-3 & Regional Dictionary ---
COUNTRY_INFO = {
    "United States": ("USA", "Americas"), "USA": ("USA", "Americas"), "US": ("USA", "Americas"), "U.S.": ("USA", "Americas"),
    "Canada": ("CAN", "Americas"), "Mexico": ("MEX", "Americas"), "Brazil": ("BRA", "Americas"), "Argentina": ("ARG", "Americas"),
    "Chile": ("CHL", "Americas"), "Colombia": ("COL", "Americas"), "Peru": ("PER", "Americas"), "Venezuela": ("VEN", "Americas"), "Cuba": ("CUB", "Americas"),
    "United Kingdom": ("GBR", "Europe"), "UK": ("GBR", "Europe"), "U.K.": ("GBR", "Europe"), "Britain": ("GBR", "Europe"),
    "Germany": ("DEU", "Europe"), "France": ("FRA", "Europe"), "Italy": ("ITA", "Europe"), "Spain": ("ESP", "Europe"),
    "Netherlands": ("NLD", "Europe"), "Belgium": ("BEL", "Europe"), "Switzerland": ("CHE", "Europe"), "Poland": ("POL", "Europe"),
    "Sweden": ("SWE", "Europe"), "Norway": ("NOR", "Europe"), "Denmark": ("DNK", "Europe"), "Finland": ("FIN", "Europe"),
    "Ireland": ("IRL", "Europe"), "Russia": ("RUS", "Europe"), "Ukraine": ("UKR", "Europe"), "European Union": ("EU", "Europe"), "EU": ("EU", "Europe"),
    "Iran": ("IRN", "West Asia/Middle East"), "Israel": ("ISR", "West Asia/Middle East"), "Saudi Arabia": ("SAU", "West Asia/Middle East"),
    "United Arab Emirates": ("ARE", "West Asia/Middle East"), "UAE": ("ARE", "West Asia/Middle East"), "Qatar": ("QAT", "West Asia/Middle East"),
    "Oman": ("OMN", "West Asia/Middle East"), "Kuwait": ("KWT", "West Asia/Middle East"), "Bahrain": ("BHR", "West Asia/Middle East"),
    "Syria": ("SYR", "West Asia/Middle East"), "Iraq": ("IRQ", "West Asia/Middle East"), "Jordan": ("JOR", "West Asia/Middle East"),
    "Lebanon": ("LBN", "West Asia/Middle East"), "Yemen": ("YEM", "West Asia/Middle East"), "Turkey": ("TUR", "West Asia/Middle East"),
    "China": ("CHN", "Asia"), "Taiwan": ("TWN", "Asia"), "Japan": ("JPN", "Asia"), "South Korea": ("KOR", "Asia"), "North Korea": ("PRK", "Asia"),
    "India": ("IND", "Asia"), "Pakistan": ("PAK", "Asia"), "Bangladesh": ("BGD", "Asia"), "Sri Lanka": ("LKA", "Asia"),
    "Vietnam": ("VNM", "Asia"), "Malaysia": ("MYS", "Asia"), "Singapore": ("SGP", "Asia"), "Indonesia": ("IDN", "Asia"),
    "Philippines": ("PHL", "Asia"), "Thailand": ("THA", "Asia"), "Myanmar": ("MMR", "Asia"), "Cambodia": ("KHM", "Asia"),
    "South Africa": ("ZAF", "Africa"), "Egypt": ("EGY", "Africa"), "Nigeria": ("NGA", "Africa"), "Kenya": ("KEN", "Africa"),
    "Ethiopia": ("ETH", "Africa"), "Morocco": ("MAR", "Africa"), "Algeria": ("DZA", "Africa"), "Sudan": ("SDN", "Africa"),
    "Congo": ("COD", "Africa"), "Democratic Republic of the Congo": ("COD", "Africa"), "Angola": ("AGO", "Africa"), "Ghana": ("GHA", "Africa"),
    "Mali": ("MLI", "Africa"), "Niger": ("NER", "Africa"), "Chad": ("TCD", "Africa"), "Somalia": ("SOM", "Africa"),
    "Australia": ("AUS", "Oceania"), "New Zealand": ("NZL", "Oceania"), "Fiji": ("FJI", "Oceania"), "Papua New Guinea": ("PNG", "Oceania")
}

INFRASTRUCTURE_DATA = {
    "Semiconductor Fabs": [
        {"name": "TSMC - Gigafab 12 (Hsinchu, Taiwan)", "lat": 24.773, "lon": 121.011},
        {"name": "TSMC - Gigafab 18 (Tainan, Taiwan)", "lat": 23.113, "lon": 120.273},
        {"name": "TSMC - JASM (Kumamoto, Japan)", "lat": 32.883, "lon": 130.866},
        {"name": "TSMC - Fab 21 (Phoenix, USA)", "lat": 33.805, "lon": -112.148},
        {"name": "Samsung - Pyeongtaek Campus (South Korea)", "lat": 37.036, "lon": 127.042},
        {"name": "Samsung - Austin Fab (USA)", "lat": 30.368, "lon": -97.625},
        {"name": "Samsung - Taylor Fab [Under Construction] (USA)", "lat": 30.565, "lon": -97.409},
        {"name": "Intel - Ocotillo Campus (Chandler, USA)", "lat": 33.262, "lon": -111.862},
        {"name": "Intel - Ronler Acres (Hillsboro, USA)", "lat": 45.542, "lon": -122.923},
        {"name": "Intel - Fab 34 (Leixlip, Ireland)", "lat": 53.374, "lon": -6.502},
        {"name": "Intel - Magdeburg [Planned] (Germany)", "lat": 52.120, "lon": 11.627},
        {"name": "SMIC - SN1/SN2 (Shanghai, China)", "lat": 31.205, "lon": 121.597},
        {"name": "SMIC - B1/B2 (Beijing, China)", "lat": 39.805, "lon": 116.505},
        {"name": "GlobalFoundries - Fab 8 (Malta, USA)", "lat": 42.970, "lon": -73.754},
        {"name": "GlobalFoundries - Fab 1 (Dresden, Germany)", "lat": 51.125, "lon": 13.714},
        {"name": "GlobalFoundries - Singapore Campus", "lat": 1.436, "lon": 103.768},
        {"name": "UMC - Fab 12A (Tainan, Taiwan)", "lat": 23.115, "lon": 120.275},
        {"name": "Micron - Boise HQ & Fab (USA)", "lat": 43.535, "lon": -116.140},
        {"name": "Micron - Hiroshima Fab (Japan)", "lat": 34.238, "lon": 132.654},
        {"name": "Texas Instruments - Sherman Campus (USA)", "lat": 33.606, "lon": -96.611},
        {"name": "Tata & PSMC (Dholera)", "lat": 22.245, "lon": 72.195},
        {"name": "CG Power (Sanand) - ACTIVE", "lat": 23.005, "lon": 72.385},
        {"name": "Micron (Sanand) - ACTIVE", "lat": 22.956, "lon": 72.338},
        {"name": "Tower Semi (Panvel) - PLANNED", "lat": 18.989, "lon": 73.117},
        {"name": "ASML HQ (Netherlands)", "lat": 51.405, "lon": 5.405},
        {"name": "HIPSPL (3DGS) - 3D Glass Packaging (Odisha, India)", "lat": 20.238, "lon": 85.702}
    ],
    "Critical Mineral Sites": [
        {"name": "Bayan Obo Mine [Largest REE] (China)", "lat": 41.796, "lon": 109.972},
        {"name": "Mountain Pass Mine [REE] (USA)", "lat": 35.473, "lon": -115.527},
        {"name": "Mount Weld Mine [REE] (Australia)", "lat": -28.775, "lon": 122.569},
        {"name": "Salar de Atacama [Lithium Triangle] (Chile)", "lat": -23.500, "lon": -68.250},
        {"name": "Mutanda Mine [Cobalt/Copper] (DRC)", "lat": -10.835, "lon": 25.795},
        {"name": "Kola Peninsula [Nickel/REE] (Russia)", "lat": 67.883, "lon": 33.000},
        {"name": "Manavalakurichi [Monazite (REEs)] (India)", "lat": 8.13, "lon": 77.30},
        {"name": "Chavara (Kollam) [REEs, Titanium] (India)", "lat": 9.01, "lon": 76.53}
    ],
    "Maritime Chokepoints": [
        {"name": "Strait of Malacca", "lat": 1.430, "lon": 103.264},
        {"name": "Strait of Hormuz", "lat": 26.566, "lon": 56.250},
        {"name": "Bab el-Mandeb (Red Sea)", "lat": 12.583, "lon": 43.333},
        {"name": "Suez Canal", "lat": 30.583, "lon": 32.333},
        {"name": "Panama Canal", "lat": 9.116, "lon": -79.750},
        {"name": "Taiwan Strait", "lat": 24.800, "lon": 119.900},
        {"name": "Bosphorus Strait", "lat": 41.221, "lon": 29.113}
    ],
    "Gulf FDI & Capital Diplomacy": [
        {"name": "Manara Minerals (Riyadh)", "lat": 24.7136, "lon": 46.6753},
        {"name": "ADQ Global Headquarters (Abu Dhabi)", "lat": 24.4539, "lon": 54.3773},
        {"name": "International Resources Holding (IRH)", "lat": 25.2048, "lon": 55.2708}
    ],
    "Naval Order of Battle & Strategic Bases": [
        {"name": "Naval Station Norfolk (USA)", "lat": 36.936, "lon": -76.326},
        {"name": "Naval Base San Diego (USA)", "lat": 32.673, "lon": -117.122},
        {"name": "Severomorsk (Russia) - Northern Fleet HQ", "lat": 69.070, "lon": 33.416},
        {"name": "US 7th Fleet HQ (Yokosuka, Japan)", "lat": 35.293, "lon": 139.661},
        {"name": "PLAN Southern Theater Command HQ (Zhanjiang, China)", "lat": 21.206, "lon": 110.402},
        {"name": "Andaman & Nicobar Command (India)", "lat": 11.666, "lon": 92.735},
        {"name": "INS Kadamba (Karwar, India)", "lat": 14.760, "lon": 74.137}
    ],
    "Aerospace & Space Force Installations": [
        {"name": "Cape Canaveral SFS / KSC (USA)", "lat": 28.488, "lon": -80.577},
        {"name": "Jiuquan Satellite Launch Center (China)", "lat": 40.960, "lon": 100.298},
        {"name": "Baikonur Cosmodrome (Kazakhstan)", "lat": 45.964, "lon": 63.305},
        {"name": "Satish Dhawan Space Centre (Sriharikota, India)", "lat": 13.719, "lon": 80.230}
    ]
}

# --- PURE PITCH BLACK CSS FIXES & MOBILE RESPONSIVENESS ---
st.markdown("""
<style>
    /* Force main app background to pure black */
    .stApp, .stAppViewContainer, .main .block-container { 
        background-color: #000000 !important; 
    }
    
    /* AGGRESSIVE FIX FOR ISSUE 2: Kill Streamlit's native crossfade/ghosting completely */
    [data-testid="stAppViewContainer"] > section > div > div,
    [data-testid="stHeader"],
    .element-container,
    .stMarkdown {
        transition: none !important;
        animation-duration: 0s !important;
        /* FIX: Removed 'opacity: 1 !important;' so the login screen can actually disappear! */
    }

    /* --- GLOBAL BUTTON THEME (Keeps Yellow Color on Dashboard & Fixes Pink Button) --- */
    div[data-testid="stButton"] > button[kind="primary"],
    div[data-testid="stButton"] > button[data-testid="baseButton-primary"],
    div[data-testid="stFormSubmitButton"] button,
    button[kind="primaryFormSubmit"] { 
        width: 100%; background-color: #facc15 !important; color: black !important; 
        font-weight: bold; border: none !important; margin-top: 10px; height: 45px;
    }
    div[data-testid="stButton"] > button[kind="primary"]:hover,
    div[data-testid="stButton"] > button[data-testid="baseButton-primary"]:hover,
    div[data-testid="stFormSubmitButton"] button:hover,
    button[kind="primaryFormSubmit"]:hover { 
        background-color: #eab308 !important; color: black !important; border: none !important;
    }
    div[data-testid="stButton"] > button[kind="primary"]:active, 
    div[data-testid="stButton"] > button[kind="primary"]:focus,
    div[data-testid="stButton"] > button[data-testid="baseButton-primary"]:active,
    div[data-testid="stFormSubmitButton"] button:active,
    button[kind="primaryFormSubmit"]:active {
        background-color: #ca8a04 !important; color: black !important; 
        border: none !important; outline: none !important; box-shadow: none !important;
    }
    
    /* Ensure Header is transparent so Mobile Hamburger remains visible on black */
    header[data-testid="stHeader"] { 
        background-color: transparent !important; 
    }
    
    /* NEW: Hide the Streamlit Running/Stop execution indicator */
    [data-testid="stStatusWidget"] {
        display: none !important;
    }

    /* NEW: Automatically handle iOS Notches and Android Punch Holes */
    .block-container {
        padding-top: max(1rem, env(safe-area-inset-top)) !important;
        padding-left: max(1rem, env(safe-area-inset-left)) !important;
        padding-right: max(1rem, env(safe-area-inset-right)) !important;
        margin-top: 0rem !important;
    }

    /* NEW: Make text adapt to small mobile screens (under 768px wide) */
    @media (max-width: 768px) {
        [data-testid="stMetricValue"] { font-size: 1.4rem !important; }
        h3 { font-size: 18px !important; }
    }
    
    [data-testid="stSidebar"] > div:first-child { padding-top: 0rem !important; }
    [data-testid="stSidebar"] { background-color: #000000 !important; border-right: 1px solid #222222 !important; } 
    [data-testid="stMetricValue"] { font-size: 1.8rem !important; font-weight: bold !important; color: #ffffff !important; }
    [data-testid="stMetricLabel"] { font-size: 1.0rem !important; white-space: normal !important; overflow: visible !important; height: auto !important; color: #aaaaaa !important; }
    footer { visibility: hidden; height: 0%; }
    h3 { margin-bottom: 0px !important; padding-bottom: 0px !important; }
    h4, h5 { margin-top: 0em !important; margin-bottom: 0.5em !important; }
    .stMarkdown p { margin-top: 5px !important; }
</style>
""", unsafe_allow_html=True)

def get_brief_mappings(directory):
    files = glob.glob(f'{directory}/brief_*.json')
    files.sort(reverse=True)
    mapping = {}
    for f in files:
        try:
            with open(f, 'r') as file:
                d = json.load(file)
                b_date = d.get('date', f.split('_')[1].split('.json')[0])
                display_name = f"SemicoN Weekly Brief - {b_date}"
                if display_name in mapping:
                    display_name = f"{display_name} (File: {f.split('_')[1].split('.json')[0]})"
                mapping[display_name] = f
        except: pass
    return mapping

def clean_dataframe(df):
    if df.empty: return df
    if 0 in df.columns: df = df.rename(columns={0: "Extracted Information"})
    if "0" in df.columns: df = df.rename(columns={"0": "Extracted Information"})
        
    cols_to_drop = [c for c in df.columns if "Unavailable" in str(c) or "None" in str(c)]
    df = df.drop(columns=cols_to_drop, errors='ignore')
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = df[col].astype(str).str.strip()
            df[col] = df[col].replace({"None": "", "N/A": "", "NA": "", "N/a": "", "n/a": "", "none": "", "nan": "", "null": ""})
    return df

def render_highlighted_text(text, keyword):
    if not text: return ""
    clean_text = text.replace('$', r'\$') 
    if keyword and keyword != "All":
        sub_keywords = [k.strip() for k in re.split(r'&|/| and |,|;', keyword) if k.strip()]
        for sub_k in sub_keywords:
            pattern = re.compile(rf'\b({re.escape(sub_k)})\b', re.IGNORECASE)
            clean_text = pattern.sub(r'<span style="background-color: #ffeb3b; color: #000000; font-weight: bold; padding: 2px 4px; border-radius: 3px;">\1</span>', clean_text)
        st.markdown(clean_text, unsafe_allow_html=True)
    else:
        st.markdown(clean_text)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000EE')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None: tblPr.remove(tblBorders)
    new_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        new_borders.append(border)
    tblPr.append(new_borders)

def add_word_data_table(doc, heading_text, data_list):
    if not data_list: return
    safe_list = []
    for item in data_list:
        if isinstance(item, dict):
            safe_list.append(item)
        else:
            safe_list.append({"Extracted Information": str(item)})
            
    if not safe_list: return
    if len(safe_list) == 1 and "No " in str(list(safe_list[0].values())[0]): return
    
    p_heading = doc.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_heading.paragraph_format.line_spacing = 1.5
    p_heading.paragraph_format.space_before = Pt(6) 
    p_heading.paragraph_format.space_after = Pt(0) 
    r_head = p_heading.add_run(heading_text)
    r_head.bold = True
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(12)

    headers = list(safe_list[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER 
    table.style = 'Table Grid'
    set_table_borders(table) 
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r.bold = True
                
    for item in safe_list:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            row_cells[i].text = str(item.get(h, ''))
            for p in row_cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
    doc.add_paragraph() 

@st.cache_data
def create_landscape_word(text_sections, final_text, actions_data, brief_date, fund_data, market_data, risk_data, sources_data, text_ews):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    header = section.header
    header.paragraphs[0].text = "" 
    r_head = header.paragraphs[0].add_run("Kashif Anwar, SemicoN Dashboard Brief")
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(11)
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer = section.footer
    footer.paragraphs[0].text = ""
    r_foot = footer.paragraphs[0].add_run("Kashif Anwar, SemicoN Dashboard Brief")
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(11)
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mast_table = doc.add_table(rows=1, cols=2)
    mast_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_logo = mast_table.cell(0, 0)
    cell_logo.width = Inches(2.0)
    cell_text = mast_table.cell(0, 1)
    cell_text.width = Inches(6.0)
    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img = Image.open("logo.jpg").convert("RGBA")
        min_dim = min(img.size)
        img = img.crop((0, 0, min_dim, min_dim))
        mask = Image.new('L', img.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, min_dim, min_dim), fill=255)
        result = Image.new('RGBA', img.size, (255, 255, 255, 0))
        result.paste(img, mask=mask)
        buf = BytesIO()
        result.save(buf, format="PNG")
        buf.seek(0)
        p_logo.add_run().add_picture(buf, width=Inches(1.5))
    except Exception: pass 
        
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_text = p_text.add_run("SemicoN Dashboard – A Semicon News Dashboard")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(20)
    run_text.bold = True
    run_text.font.color.rgb = RGBColor(255, 0, 0) 
    doc.add_paragraph() 
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.5
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12) 
    run_title = p_title.add_run(f'SemicoN Weekly Brief - {brief_date}')
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.name = 'Times New Roman'
    
    p_author = doc.add_paragraph()
    p_author.paragraph_format.line_spacing = 1.5
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after = Pt(12)
    run_author = p_author.add_run("Prepared By: Kashif Anwar")
    run_author.bold = True
    run_author.font.size = Pt(12)
    run_author.font.name = 'Times New Roman'
    
    section_titles = [
        "Executive Summary",
        "Global Foundry Market & Geopolitical Positioning",
        "AI Chip Demand, Manufacturing & Processing",
        "Critical Minerals: Rare Earth Reserves & Supply Chains",
        "Export Controls & Geopolitical Impact",
        "AI, Chips and Rare Earth in Military and Outer Space Domain",
        "Lithography Chokepoints & State Actions",
        "India: Domestic & Strategic Developments",
        "West Asia/Middle East: Domestic & Strategic Developments"
    ]

    for i, text_data in enumerate(text_sections):
        
        if i == 0 and text_ews and text_ews.strip() != "":
            p_ews_head = doc.add_paragraph()
            p_ews_head.paragraph_format.space_before = Pt(12)
            r_ews_head = p_ews_head.add_run("🚨 Early Warning & Red Flags")
            r_ews_head.bold = True
            r_ews_head.font.color.rgb = RGBColor(255, 0, 0)
            r_ews_head.font.name = 'Times New Roman'
            r_ews_head.font.size = Pt(14)
            
            p_ews = doc.add_paragraph()
            r_ews = p_ews.add_run(text_ews)
            r_ews.font.name = 'Times New Roman'
            r_ews.font.size = Pt(12)
            
        if not text_data or text_data.strip() == "": continue

        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(18)
        r_head = p_head.add_run(section_titles[i])
        r_head.bold = True
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.color.rgb = RGBColor(0, 102, 204)

        for line in text_data.split('\n'):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            
            if line.startswith('**') and line.endswith('**'):
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.replace('**', ''))
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif line.startswith('* '):
                p.style = 'List Bullet'
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line[2:].replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line.replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if i == 1: add_word_data_table(doc, "Strategic Investments & Funding", fund_data)
        if i == 3: add_word_data_table(doc, "Market & Geopolitical Impact", market_data)
        if i == 4: add_word_data_table(doc, "Supply Chain Risk Analysis", risk_data)
                
    add_word_data_table(doc, 'Recent State Actions', actions_data)

    if final_text and final_text.strip() != "":
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(18)
        r_head = p_head.add_run("Strategic Conclusion")
        r_head.bold = True
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(14)
        r_head.font.color.rgb = RGBColor(0, 102, 204)

        for line in final_text.split('\n'):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            
            if line.startswith('**') and line.endswith('**'):
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line.replace('**', ''))
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif line.startswith('* '):
                p.style = 'List Bullet'
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line[2:].replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line.replace('**', ''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
    if sources_data:
        doc.add_paragraph() 
        p_source_head = doc.add_paragraph()
        p_source_head.paragraph_format.line_spacing = 1.5
        p_source_head.paragraph_format.space_before = Pt(12)
        r_source = p_source_head.add_run("Verified Intelligence Sources")
        r_source.bold = True
        r_source.font.name = 'Times New Roman'
        r_source.font.size = Pt(14)
        
        for src in sources_data:
            p_src = doc.add_paragraph()
            p_src.style = 'List Bullet'
            p_src.paragraph_format.space_after = Pt(6)
            p_src.paragraph_format.line_spacing = 1.0
            add_hyperlink(p_src, src['url'], src['title'])
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_latest_file():
    if not os.path.exists('data'): return None
    files = glob.glob('data/brief_*.json')
    if not files: return None
    files.sort() 
    return files[-1]

latest_filepath = get_latest_file()

@st.cache_data(ttl=30) 
def load_data(filepath):
    if not filepath: return None
    with open(filepath, 'r') as f: return json.load(f)

dashboard_data = load_data(latest_filepath)

# ==========================================
# TEXT PARSER FOR RSS ACCUMULATOR FIX
# ==========================================
def parse_rss_txt_file():
    import urllib.parse
    from datetime import datetime, timedelta, timezone

    rss_dict = {}
    filepath = 'data/rss_accumulator.txt'
    if not os.path.exists(filepath): return rss_dict

    # --- CRITICAL FIX: FREEZE 24-HOUR WINDOW TO 00:15 AM IST ROLLOVER ---
    now_utc = datetime.now(timezone.utc)
    now_ist = now_utc + timedelta(hours=5, minutes=30)

    # The new target anchor is 18:45 UTC (00:15 AM IST).
    # We freeze the dashboard to ONLY evaluate news from the 24 hours PRECEDING this exact time.
    if now_ist.hour == 0 and now_ist.minute < 15:
        # If visited before 00:15 AM IST, lock onto yesterday's rollover snapshot
        anchor_ist = now_ist.replace(hour=0, minute=15, second=0, microsecond=0) - timedelta(days=1)
    else:
        # If visited after 00:15 AM IST, lock onto today's rollover snapshot
        anchor_ist = now_ist.replace(hour=0, minute=15, second=0, microsecond=0)

    # Convert strict window back to UTC for safe comparison with published timestamps
    window_end_utc = anchor_ist - timedelta(hours=5, minutes=30)
    window_start_utc = window_end_utc - timedelta(hours=24)
    # ----------------------------------------------------------------------

    current_reg = None

    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line.startswith("---") and "DOMAIN-FORCED NEWS" in line:
                reg = line.replace("---", "").replace("DOMAIN-FORCED NEWS", "").strip()
                if "Middle East" in reg: reg = "West Asia/Middle East"
                current_reg = reg
                if current_reg not in rss_dict:
                    rss_dict[current_reg] = []
            elif line.startswith("- [") and current_reg:
                try:
                    d_end = line.find("]")
                    date_str = line[3:d_end]
                    title_str = line[d_end+1:].strip()

                    # --- NEW STATIC WINDOW EVALUATION ---
                    is_recent = False
                    if date_str != "Recent Update":
                        try:
                            pub_dt = datetime.strptime(date_str, "%Y-%m-%d %H:%M").replace(tzinfo=timezone.utc)
                            # ONLY mark as 24h if it falls perfectly inside the locked daily window
                            if window_start_utc <= pub_dt <= window_end_utc:
                                is_recent = True
                        except Exception:
                            is_recent = False
                    else:
                        is_recent = False

                    clean_search = title_str.replace("🔴", "").replace("🟠", "").replace("🟡", "").replace("CRITICAL:", "").replace("ELEVATED:", "").replace("WATCH:", "").replace("LIVE WARNING:", "").strip()
                    search_query = urllib.parse.quote_plus(clean_search)
                    news_link = f"https://news.google.com/search?q={search_query}"

                    if not any(x['title'] == title_str for x in rss_dict[current_reg]):
                        # Inject the "is_24h" flag into the dictionary
                        rss_dict[current_reg].append({"title": title_str, "published": date_str, "link": news_link, "is_24h": is_recent})
                except Exception: pass
    return rss_dict

def extract_tag(tag, text):
    match = re.search(rf'<{tag}>(.*?)</{tag}>', text, re.DOTALL | re.IGNORECASE)
    if match:
        clean_text = match.group(1).strip()
        clean_text = re.sub(r'</?[A-Z_]+>', '', clean_text)
        return clean_text
    return None

if dashboard_data:
    brief_date = dashboard_data.get('date', 'Latest Integration')
    raw_text = dashboard_data.get('brief_raw', '')
    t_sum = extract_tag('SUMMARY', raw_text)
    t_ews = extract_tag('EWS', raw_text)
    t1 = extract_tag('EXEC', raw_text)
    t2 = extract_tag('LITHO', raw_text)
    t3 = extract_tag('REE', raw_text)
    t4 = extract_tag('GEO', raw_text)
    t_mil = extract_tag('MILITARY', raw_text)
    t5 = extract_tag('CONCLUSION', raw_text)
    t_india = extract_tag('INDIA', raw_text)
    t_wa = extract_tag('WEST_ASIA', raw_text)
    t_fin = extract_tag('FINAL_CONCLUSION', raw_text)
    
    text_summary = t_sum or ""
    text_ews = t_ews or ""
    text_section_1 = t1 or ""
    text_section_2 = t2 or ""
    text_section_3 = t3 or ""
    text_section_4 = t4 or ""
    text_military = t_mil or ""
    text_section_5 = t5 or ""
    text_india = t_india or ""
    text_wa = t_wa or ""
    text_final = t_fin or ""
else:
    brief_date = "System Offline"
    raw_text = ""
    text_summary = text_ews = text_section_1 = text_section_2 = text_section_3 = text_section_4 = text_military = text_section_5 = text_india = text_wa = text_final = "Awaiting deployment of new intelligence brief."
    dashboard_data = {"supply_chain_risk": [{"Risk Factor": "No data"}], "recent_actions": [], "funding_data": [{"Entity": "No data"}], "market_impact": [{"Entity": "No data"}]}

# ==========================================
# NEW ALGORITHMIC THREAT SCORING FUNCTION
# ==========================================
def calculate_domain_threat(domain_name, text_content, dash_data):
    """
    Dynamically calculates a threat score (0-100) based on textual severity 
    and cross-referencing json data (risks & actions) without hardcoding values.
    """
    if not text_content or len(text_content.strip()) < 20:
        return 0
    
    score = 35 

    text_lower = text_content.lower()
    critical_keywords = ['ban', 'sanction', 'shortage', 'escalation', 'military', 'war', 'blockade', 'strike', 'chokepoint', 'threat', 'breach', 'crisis']
    high_keywords = ['tariff', 'control', 'restrict', 'vulnerability', 'disrupt', 'tension', 'export control', 'embargo', 'risk']
    medium_keywords = ['delay', 'subsidy', 'compete', 'invest', 'shift', 'policy', 'regulate', 'pressure', 'concern', 'geopolitical']

    score += sum(text_lower.count(kw) for kw in critical_keywords) * 8
    score += sum(text_lower.count(kw) for kw in high_keywords) * 5
    score += sum(text_lower.count(kw) for kw in medium_keywords) * 2

    domain_parts = [p.lower() for p in domain_name.replace(" / ", " ").split() if len(p) > 3]

    risks = dash_data.get('supply_chain_risk', [])
    for risk in risks:
        risk_text = str(risk).lower()
        if any(part in risk_text for part in domain_parts):
            score += 10 
    
    actions = dash_data.get('recent_actions', [])
    for action in actions:
        action_text = str(action).lower()
        if any(part in action_text for part in domain_parts):
            score += 5 

    return min(100, max(20, score))


# --- NEW CENTRALIZED LIVE ALERT FETCH FUNCTION ---
def get_active_live_alert():
    if not os.path.exists('data/live_alert.json'): 
        return None
    try:
        with open('data/live_alert.json', 'r') as f:
            alert = json.load(f)
        alert_time = datetime.fromisoformat(alert['timestamp'].replace("Z", "+00:00"))
        time_diff = datetime.now(timezone.utc) - alert_time
        
        if time_diff.total_seconds() < 7200: 
            return alert
    except Exception:
        pass 
    return None

# --- NEW ROCK-SOLID CLOCK CACHE (FIX FOR ISSUE 2) ---
def get_deployment_timestamp():
    """Anchors the clock to a persistent file so it survives all refreshes and log-ins."""
    os.makedirs('data', exist_ok=True)
    file_path = 'data/nominal_timer.txt'
    
    # Read the saved timestamp if it exists
    if os.path.exists(file_path):
        try:
            with open(file_path, 'r') as f:
                return int(f.read().strip())
        except Exception:
            pass
            
    # If no file exists (first run), generate the time and save it
    now_ms = int(time.time() * 1000)
    try:
        with open(file_path, 'w') as f:
            f.write(str(now_ms))
    except Exception:
        pass
        
    return now_ms

def check_early_warnings():
    try:
        if st.session_state.get('role') == 'admin' and os.path.exists('data/live_alert.json'):
            if st.button("🛠️ Admin: Force Clear Live Alert (Resolve Stuck EWS)", type="primary"):
                try:
                    os.remove('data/live_alert.json')
                    st.success("Alert cleared! Refreshing...")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"Failed to clear: {e}")

        alert = get_active_live_alert()
        
        box_bg_color = "#000000"
        nominal_text_color = "#d1d5db"
        
        # Guarantee the fallback time never shifts during navigation or refresh
        fallback_time_ms = get_deployment_timestamp()
        
        if alert:   
            try:
                # Safely check for timestamp without triggering datetime.now() execution
                if 'timestamp' in alert:
                    dt = datetime.fromisoformat(alert['timestamp'].replace("Z", "+00:00"))
                    start_timestamp_ms = int(dt.timestamp() * 1000)
                else:
                    start_timestamp_ms = fallback_time_ms
            except:
                start_timestamp_ms = fallback_time_ms
            
            # Pre-calculate the elapsed time in Python to eliminate the 00:00:00 visual flash
            elapsed_ms = max(0, int(time.time() * 1000) - start_timestamp_ms)
            h, m, s = elapsed_ms // 3600000, (elapsed_ms % 3600000) // 60000, (elapsed_ms % 60000) // 1000
            initial_clock = f"{h:02d}:{m:02d}:{s:02d}"
            
            # --- RESPONSIVE DEFCON CSS FIX ---
            html_code = f"""
            <style>
                body {{ font-family: 'Courier New', Courier, monospace; margin: 0; padding: 0; background-color: {box_bg_color}; overflow: hidden; }}
                .defcon-box {{
                    background: linear-gradient(90deg, #8b0000 0%, #ff0000 50%, #8b0000 100%); 
                    background-size: 200% 200%; 
                    animation: pulseBackground 2s infinite; 
                    border: 2px solid #ff4b4b; 
                    padding: 15px; 
                    border-radius: 8px; 
                    box-shadow: 0 0 15px rgba(255, 0, 0, 0.5);
                    color: #ffffff;
                    min-height: 185px;
                    max-height: 240px; /* NEW: Forces the box to stop expanding and trigger the scrollbar */
                    height: auto;
                    box-sizing: border-box;
                    overflow-y: auto; 
                    -webkit-overflow-scrolling: touch; /* NEW: Forces smooth scroll support on Android/iOS non-Safari browsers */
                }}
                :fullscreen {{
                    background-color: rgba(20, 20, 20, 0.95);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }}
                :fullscreen .defcon-box {{ width: 90vw; max-height: 90vh; height: auto; padding: 40px; border-width: 4px; }}
                :fullscreen .title {{ font-size: 2.5em; }}
                :fullscreen .timer {{ font-size: 1.5em; }}
                :fullscreen .headline {{ font-size: 2em; line-height: 1.2; margin-top: 20px; }}
                :fullscreen .summary {{ font-size: 1.5em; line-height: 1.4; margin-top: 20px; }}
                
                @keyframes pulseBackground {{ 0% {{ background-position: 0% 50%; }} 50% {{ background-position: 100% 50%; }} 100% {{ background-position: 0% 50%; }} }}
                @keyframes blinkText {{ 0%, 100% {{ opacity: 1; }} 50% {{ opacity: 0; }} }}
                
                /* Responsive Flex Wrap Fix */
                .header-flex {{ display: flex; justify-content: space-between; align-items: center; margin-top: 0px; margin-bottom: 8px; flex-wrap: wrap; gap: 10px; }}
                .title {{ font-size: 1.17em; font-weight: bold; letter-spacing: 2px; text-transform: uppercase; margin:0; display:flex; align-items:center; flex: 1 1 100%; }}
                .timer-container {{ display: flex; align-items: center; flex: 1 1 100%; justify-content: flex-start; margin-bottom: 5px; }}
                
                @media (min-width: 600px) {{
                    .title {{ flex: 1; }}
                    .timer-container {{ flex: 1; justify-content: flex-end; margin-bottom: 0px; }}
                }}
                
                .timer {{ font-size: 15px; font-weight: bold; background: rgba(0,0,0,0.5); padding: 5px 10px; border-radius: 5px; border: 1px solid rgba(255,255,255,0.4); }}
                .headline {{ font-weight: 800; font-size: 16px; margin-bottom: 8px; margin-top:0; }}
                .summary {{ font-size: 14px; color: #f8f8f8; margin-bottom: 0px; border-top: 1px solid rgba(255,255,255,0.3); padding-top: 10px; }}
                
                .magnify-btn {{ background: rgba(0,0,0,0.6); color: white; border: 1px solid white; padding: 4px 8px; border-radius: 4px; cursor: pointer; font-size: 12px; font-weight: bold; margin-left: 10px; }}
                .magnify-btn:hover {{ background: rgba(255,255,255,0.2); }}
                :fullscreen .magnify-btn {{ display: none; }} 
                .close-btn {{ display: none; }}
                :fullscreen .close-btn {{ display: inline-block; background: transparent; color: white; border: 1px solid white; padding: 8px 16px; border-radius: 4px; cursor: pointer; font-size: 18px; margin-top: 20px; }}
            </style>
            
            <div class="defcon-box" id="defcon-container">
                <div class="header-flex">
                    <h3 class="title">
                        <span style="animation: blinkText 1s infinite; margin-right: 15px;">⚠️ DEFCON-LEVEL THREAT</span> 
                    </h3>
                    <div class="timer-container">
                        <div class="timer">Live Since: <span id="clock">{initial_clock}</span></div>
                        <button class="magnify-btn" onclick="toggleFullscreen()">🔍 MAGNIFY</button>
                    </div>
                </div>
                <p class="headline">{alert.get('headline', '')}</p>
                <p class="summary">{alert.get('summary', '')}</p>
                <button class="close-btn" onclick="document.exitFullscreen()">✖ CLOSE VIEW</button>
            </div>
            
            <script>
                function toggleFullscreen() {{
                    let elem = document.documentElement;
                    if (!document.fullscreenElement) {{
                        elem.requestFullscreen().catch(err => {{ alert(`Error: ${{err.message}}`); }});
                    }} else {{ document.exitFullscreen(); }}
                }}
                const start = {start_timestamp_ms};
                function update() {{
                    const now = new Date().getTime();
                    let diff = now - start;
                    if(diff < 0) diff = 0;
                    let h = Math.floor(diff / (1000 * 60 * 60));
                    let m = Math.floor((diff % (1000 * 60 * 60)) / (1000 * 60));
                    let s = Math.floor((diff % (1000 * 60)) / 1000);
                    document.getElementById('clock').innerText = String(h).padStart(2, '0') + ':' + String(m).padStart(2, '0') + ':' + String(s).padStart(2, '0');
                }}
                setInterval(update, 1000); update();
            </script>
            """
            components.html(html_code, height=260) 
            
        else:
            start_timestamp_ms = fallback_time_ms
            
            # Pre-calculate to eliminate 00:00:00 flash
            elapsed_ms = max(0, int(time.time() * 1000) - start_timestamp_ms)
            h, m, s = elapsed_ms // 3600000, (elapsed_ms % 3600000) // 60000, (elapsed_ms % 60000) // 1000
            initial_clock = f"{h:02d}:{m:02d}:{s:02d}"

            # --- RESPONSIVE NOMINAL CSS FIX ---
            html_code = f"""
            <style>
                body {{ font-family: system-ui, -apple-system, sans-serif; margin: 0; padding: 0; background-color: {box_bg_color}; overflow: hidden; }}
                .nominal-box {{ background-color: rgba(0, 191, 255, 0.05); border-left: 5px solid #00bfff; padding: 15px; border-radius: 5px; min-height: 90px; height: auto; }}
                .header-flex {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px; flex-wrap: wrap; gap: 8px; }}
                .title {{ color: #00bfff; margin: 0; font-size: 1.1em; font-weight: bold; display: flex; align-items: center; }}
                .timer {{ font-size: 13px; font-weight: bold; color: #00bfff; background: rgba(0, 191, 255, 0.1); padding: 5px 10px; border-radius: 4px; border: 1px solid rgba(0, 191, 255, 0.3); font-family: monospace; }}
                .desc {{ font-size: 14px; margin: 0; color: {nominal_text_color}; }}
            </style>
            <div class="nominal-box">
                <div class="header-flex">
                    <h4 class="title"><span>🟢 System Nominal</span></h4>
                    <div class="timer">Status Verified: <span id="clock">{initial_clock}</span> ago</div>
                </div>
                <p class="desc">Current Warning System doesn't see an early warning situation. Watch out for further updates.</p>
            </div>
            <script>
                const start = {start_timestamp_ms};
                function update() {{
                    const now = new Date().getTime();
                    let diff = now - start;
                    if(diff < 0) diff = 0;
                    let h = Math.floor(diff / (1000 * 60 * 60));
                    let m = Math.floor((diff % (1000 * 60 * 60)) / (1000 * 60));
                    let s = Math.floor((diff % (1000 * 60)) / 1000);
                    document.getElementById('clock').innerText = String(h).padStart(2, '0') + ':' + String(m).padStart(2, '0') + ':' + String(s).padStart(2, '0');
                }}
                setInterval(update, 1000); update();
            </script>
            """
            components.html(html_code, height=130)

    except Exception as e:
        pass  # <--- THIS CLOSES THE TRY BLOCK AT THE TOP OF THE FUNCTION

# ==========================================
# SHOCKWAVE ENGINE V2 (INTELLIGENCE GRADE)
# ==========================================
def run_shockwave_engine():
    from datetime import datetime, timedelta, timezone
    import pandas as pd
    import streamlit as st
    import re

    st.markdown("<h3 style='color:#ff9f1c;'>🌍 Geopolitical Shockwave Engine - In the Last 24 Hours</h3>", unsafe_allow_html=True)

    try:
        rss_data = parse_rss_txt_file()
        if not rss_data:
            st.warning("No RSS data available.")
            return

        # ---- Flatten last 24h news ----
        news_items = []
        for region, articles in rss_data.items():
            for art in articles:
                if art.get("is_24h", False):  # Ensure we strictly pull the 24h flagged items
                    news_items.append(art)

        if not news_items:
            st.info("No events detected in last 24h.")
            return

        # ---- Domain Keywords ----
        domain_map = {
            "Supply Chain": ["shortage","delay","logistics","disruption","shutdown"],
            "Rare Earths": ["rare earth","mineral","lithium","cobalt","mining"],
            "AI Chips": ["ai chip","nvidia","gpu","semiconductor","tsmc"],
            "Export Controls": ["ban","sanction","export control","restriction","embargo"],
            "Military Tech": ["military","defense","missile","war","navy","air force"]
        }

        # ---- Impact weights ----
        propagation_weights = {
            "Supply Chain": {"Rare Earths":0.6,"AI Chips":0.8,"Export Controls":0.5},
            "Rare Earths": {"Supply Chain":0.7,"AI Chips":0.6},
            "AI Chips": {"Supply Chain":0.5,"Military Tech":0.7},
            "Export Controls": {"AI Chips":0.9,"Supply Chain":0.6},
            "Military Tech": {"AI Chips":0.8,"Supply Chain":0.4}
        }

        shock_events = []

        for item in news_items:
            title = item.get("title","")
            title_lower = title.lower()

            base_score = 0
            origin_domain = None

            # ---- Detect origin ----
            for domain, keywords in domain_map.items():
                if any(kw in title_lower for kw in keywords):
                    origin_domain = domain
                    base_score += 5

            # Extra severity boost
            if any(k in title_lower for k in ["war","crisis","ban","sanction","strike"]):
                base_score += 5

            if not origin_domain:
                continue

            base_score = min(base_score, 10)

            # ---- Propagation ----
            propagation_result = {}

            if origin_domain in propagation_weights:
                for target, weight in propagation_weights[origin_domain].items():
                    propagation_result[target] = round(base_score * weight, 2)

            total_impact = base_score + sum(propagation_result.values())

            shock_events.append({
                "Event": title[:80],
                "Origin": origin_domain,
                "Base Impact": base_score,
                "Ripple Impact": round(sum(propagation_result.values()),2),
                "Total Shock Score": round(total_impact,2)
            })

        if not shock_events:
            st.info("No major shockwaves detected.")
            return

        df = pd.DataFrame(shock_events)
        df = df.sort_values(by="Total Shock Score", ascending=False).head(10)

        # ---- Global Risk Index ----
        global_score = int(df["Total Shock Score"].mean())

        if global_score > 20:
            status = "🔴 Critical"
        elif global_score > 12:
            status = "🟠 Elevated"
        else:
            status = "🟡 Watch"

        # ---- UI OUTPUT ----
        col1, col2 = st.columns([2,1])

        with col1:
            st.markdown("#### Top Shockwave Events (Last 24H)")
            st.dataframe(df, use_container_width=True)

        with col2:
            st.markdown("#### Global Shock Index")
            st.metric("Score", f"{global_score}/30", status)

        # ---- Intelligence Summary ----
        top_event = df.iloc[0]

        st.markdown("#### 🧠 Intelligence Assessment")
        st.markdown(f"""
        **Primary Shock Driver:** {top_event['Event']}  
        **Origin Domain:** {top_event['Origin']}  

        **Analysis:**
        This event is propagating across multiple semiconductor dependencies.  
        Secondary disruptions are visible in adjacent domains due to structural interdependence.

        **System Impact Pathway:**
        {top_event['Origin']} → Supply Chain → AI Chips → Strategic Risk Escalation

        **Assessment:**
        Current environment reflects **{status} geopolitical volatility** with cascading supply chain implications.
        """)

    except Exception as e:
        st.error(f"Shockwave Engine Error: {e}")

# ==========================================
# 2. TICKER TAPE & CSS INJECTIONS
# ==========================================
def render_ticker_tape():
    ticker_items = []
    
    try:
        # Automatically grab the data
        live_rss = parse_rss_txt_file()
        if not live_rss: return
        
        # Flatten the region dictionary into a single list of unique news items
        seen_titles = set()
        unique_news = []
        for region, articles in live_rss.items():
            for art in articles:
                if art['title'] not in seen_titles and art.get('is_24h', False): # Only 24h news
                    seen_titles.add(art['title'])
                    unique_news.append(art)

        # Dynamic Threat Scoring ...
        critical = ['ban', 'sanction', 'shortage', 'escalation', 'military', 'war', 'blockade', 'strike', 'chokepoint', 'threat', 'breach', 'crisis']
        high = ['tariff', 'control', 'restrict', 'vulnerability', 'disrupt', 'tension', 'export control', 'embargo', 'risk']
        med = ['delay', 'subsidy', 'compete', 'invest', 'shift', 'policy', 'regulate', 'pressure', 'concern', 'geopolitical']

        for item in unique_news:
            title_lower = item['title'].lower()
            score = 3
            score += sum(1 for kw in critical if kw in title_lower) * 5
            score += sum(1 for kw in high if kw in title_lower) * 3
            score += sum(1 for kw in med if kw in title_lower) * 2
            score = min(10, score)
            
            if score >= 5:
                # Keep the emojis for visual context, but drop the color assignments
                if score >= 9:
                    prefix = "🔴 CRITICAL:"
                elif score >= 7:
                    prefix = "🟠 ELEVATED:"
                else:
                    prefix = "🟡 WATCH:"
                
                clean_title = item.get("title", "").replace('"', '&quot;').replace("'", "&#39;")
                ticker_html = f'<div class="ticker-item"><a href="{item.get("link", "#")}" target="_blank">{prefix} {clean_title}</a></div>'
                ticker_items.append(ticker_html)
                
    except Exception as e:
        pass

    if not ticker_items: return
    all_items_html = "".join(ticker_items)

    # --- DYNAMIC SPEED CALCULATION ---
    # Assign ~10 seconds per news item, plus 15 seconds to travel across the screen.
    # This guarantees consistent reading speed regardless of how much news there is!
    dynamic_duration = max(20, len(ticker_items) * 10 + 15)

    # Inject CSS
    ticker_code = f"""
    <style>
    /* ================================
    SEMICON INTELLIGENCE TICKER
    ================================ */

    .ticker-wrap {{
        position: fixed;
        top: 60px;
        left: 0;
        width: 100vw;
        height: 42px;
        background-color: #050505;
        border: none !important; /* Completely removes the outline */
        box-shadow: none !important; /* Removes any bottom shadow */
        z-index: 990; 
        overflow: hidden;
        display: flex;
        align-items: center;
    }}

    .block-container {{
        padding-top: 110px !important; 
    }}

    .ticker-move {{
        display: inline-block;
        white-space: nowrap;
        padding-left: 100vw;
        /* Inject the Python calculated duration dynamically */
        animation: ticker {dynamic_duration}s linear infinite;
    }}

    .ticker-move:hover {{
        animation-play-state: paused;
    }}

    @keyframes ticker {{
        0% {{ transform: translate3d(0,0,0); }}
        100% {{ transform: translate3d(-100%,0,0); }}
    }}

    .ticker-item {{
        display: inline-block;
        margin-right: 60px;
        font-family: "Courier New", monospace;
        font-weight: bold;
        font-size: 14px;
        letter-spacing: 0.5px;
    }}

    .ticker-item a {{
        text-decoration: none;
        color: #ffffff !important; /* Forces all text to be pure white */
    }}

    .ticker-item a:hover {{
        text-decoration: underline;
        opacity: 0.8;
    }}
    </style>

    <div class="ticker-wrap">
        <div class="ticker-move">
            {all_items_html}
        </div>
    </div>
    """
    st.markdown(ticker_code, unsafe_allow_html=True)

# ==========================================
# LOGIN SCREEN 
# ==========================================
if 'role' not in st.session_state: st.session_state['role'] = None

if st.session_state['role'] is None:
    
    # --- FIX: Wrap login in a placeholder to instantly destroy it upon authentication ---
    login_placeholder = st.empty()
    
    with login_placeholder.container():
        st.markdown("""
        <div class="fixed-left-panel">
            <h1>Be a Part of<br>Something <span>Beautiful</span></h1>
            <p>Access high-fidelity insights at the intersection of global policy and the semiconductor industry.</p>
        </div>
        """, unsafe_allow_html=True)

        # Safely convert logo to base64 so it securely renders inside HTML
        try:
            with open("logo.jpg", "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode()
                img_html = f'<img src="data:image/jpeg;base64,{encoded_string}" class="login-logo" style="width: 300px; max-width: 100%;"/>'
        except FileNotFoundError:
            img_html = '' 

        # Inject the Single Unified Block ensuring the logo and text are permanently attached
        st.markdown(f"""
        <div class="login-header">
            {img_html}
            <h2>Login</h2>
            <p>Enter your credentials</p>
        </div>
        
        <style>
            /* Hack to remove form border natively across all Streamlit versions */
            [data-testid="stForm"] {{ border: none !important; padding: 0 !important; }}
        </style>
        """, unsafe_allow_html=True)
        
        # --- FIX 1: Use columns to shrink the width of the login elements ---
        spacer_left, center_col, spacer_right = st.columns([1, 1.5, 1])
        
        with center_col:
            # --- FIX 2: Wrap inside st.form to stop keystroke & eye-icon flickering ---
            with st.form("login_form", border=False):
                email_input = st.text_input("Email", placeholder="analyst@agency.gov")
                password_input = st.text_input("Password", type="password", placeholder="Enter Secure Key")
                
                # Note: st.button becomes st.form_submit_button inside a form
                submit_login = st.form_submit_button("Login", type="primary", use_container_width=True)

            if submit_login:
                ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "anwarkashif@semirare.in")
                ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "Admin@123") 
                
                if email_input == ADMIN_EMAIL and password_input == ADMIN_PASSWORD:
                    st.session_state['role'] = 'admin'
                    login_placeholder.empty() # INSTANTLY DESTROY LOGIN WIDGETS BEFORE RERUN
                    st.rerun()
                else: 
                    st.toast("Invalid credentials. Please verify your secure key.", icon="🚫")

            # --- GUEST ACCESS ---
            # Kept inside the center_col so it perfectly matches the constrained width
            if st.button("View as Guest", type="secondary", use_container_width=True):
                st.session_state['role'] = 'guest'
                login_placeholder.empty() # INSTANTLY DESTROY LOGIN WIDGETS BEFORE RERUN
                st.rerun()

# ==========================================
# 4. MAIN DASHBOARD
# ==========================================
else:
    # --- FIX: USE SESSION STATE SO SPLASH SCREEN ONLY PLAYS ONCE ---
    if 'splash_shown' not in st.session_state:
        st.session_state['splash_shown'] = True # Marks that it has played
        
        st.markdown("""
        <style>
            .splash-overlay {
                position: fixed;
                top: 0;
                left: 0;
                width: 100vw;
                height: 100vh;
                background-color: #000000; /* Pitch dark black */
                display: flex;
                align-items: center;
                justify-content: center;
                z-index: 9999999; /* Forces it on top of EVERYTHING */
                animation: fadeOutSplash 3.5s forwards; 
                pointer-events: none; /* Allows user to interact with dashboard once it fades */
            }
            .splash-text {
                color: #ffffff; /* White text */
                font-size: 2.5rem;
                font-weight: 300;
                font-family: 'Times New Roman', Times, serif;
                letter-spacing: 2px;
                text-align: center;
            }
            @keyframes fadeOutSplash {
                0% { opacity: 1; visibility: visible; }
                75% { opacity: 1; visibility: visible; } 
                100% { opacity: 0; visibility: hidden; display: none; }
            }
            @keyframes blinkDots {
                0%, 100% { opacity: 0; }
                50% { opacity: 1; }
            }
            .loading-dots {
                animation: blinkDots 1.2s infinite ease-in-out;
            }
        </style>
        
        <div class="splash-overlay">
            <div class="splash-text">Welcome to my SemicoN Dashboard<span class="loading-dots">...</span></div>
        </div>
        """, unsafe_allow_html=True)

    # --- RENDER THE NEW TICKER TAPE FIRST ---
    # (The CSS inside this function safely positions the native Streamlit toggle)
    render_ticker_tape()

    # --- SIDEBAR CONTENT ---
    # Because we hijacked the native button, Streamlit manages opening/closing natively!
    st.sidebar.image("logo.jpg", use_container_width=True)
    st.sidebar.markdown("""
    <div style='text-align: center; margin-top: -10px;'>
        <p style='font-size: 16px; font-weight: bold; margin-bottom: 5px;'>A Semicon News Dashboard.</p>
        <p style='font-size: 14px; margin-bottom: 5px;'>Prepared by: Kashif Anwar</p>
        <p style='font-size: 13px; color: #00bfff; font-weight: bold; font-style: italic; margin-bottom: 0px;'>A Human-AI Vetted Dashboard</p>
    </div>
    """, unsafe_allow_html=True)
    st.sidebar.markdown("---")
    
    st.sidebar.markdown("<p style='font-size: 16px; font-weight: bold; margin-bottom: 5px;'>Global Filter</p>", unsafe_allow_html=True)
    
    actor_list = []
    
    df_actions = clean_dataframe(pd.DataFrame(dashboard_data.get('recent_actions', [])))
    
    if not df_actions.empty and 'Actor' in df_actions.columns:
        actor_list = [a for a in df_actions['Actor'].dropna().unique().tolist() if str(a).strip()]
    
    selected_actor = st.sidebar.selectbox("🔍 Highlight & Filter by Actor:", ["All"] + sorted(actor_list))
    st.sidebar.markdown("---")

    st.sidebar.markdown("<p style='font-size: 16px; font-weight: bold; margin-bottom: 5px;'>Advanced Threat Tools</p>", unsafe_allow_html=True)
    advanced_tool = st.sidebar.radio(
        "Launch Sandbox Module:", 
        ["None (View Main Dashboard)", "Quantitative Threat Scoring", "Intelligence Interrogation (RAG)"], 
        index=0
    )
    st.sidebar.markdown("---")

    st.sidebar.markdown("<p style='font-size: 16px; font-weight: bold; margin-bottom: 5px;'>Regions Covered</p>", unsafe_allow_html=True)
    
    region_checks = {
        "Asia": ["**Asia:**", text_india], 
        "Middle East/West Asia": ["**Middle East/West Asia:**", text_wa],
        "Africa": ["**Africa:**"],
        "Europe": ["**Europe:**"],
        "Americas": ["**Americas:**"],
        "Oceania": ["**Oceania:**"]
    }
    
    for r, indicators in region_checks.items():
        is_covered = False
        for ind in indicators:
            if ind and len(ind.strip()) > 5: 
                is_covered = True
            elif ind in raw_text: 
                is_covered = True
                
        if is_covered:
            st.sidebar.markdown(f"✅ {r}")
        else:
            st.sidebar.markdown(f"➖ <span style='color:grey'>{r}</span>", unsafe_allow_html=True)
            
    st.sidebar.markdown("---")
    with st.sidebar.expander("🧠 Key Concepts Explained"):
        st.markdown("""
        **EUV Lithography:** Extreme Ultraviolet Lithography. The cutting-edge tech used to print the most advanced microchips.  
        **Tape-out:** The final design phase for integrated circuits before manufacturing begins.  
        **Foundry:** A specialized factory where semiconductor chips are manufactured (e.g., TSMC).  
        **Rare Earth Elements (REEs):** 17 metallic elements crucial for high-tech, defense, and green energy products.  
        **Fabless:** Companies that design chips but outsource manufacturing (e.g., Nvidia, Apple, AMD).
        """)

    st.sidebar.markdown("---")
    st.sidebar.title("SemicoN Access")

    if st.session_state['role'] == 'admin':
        st.sidebar.info("Access Level: **Administrator**")
        if st.sidebar.button("Logout"):
            st.session_state['role'] = None
            st.rerun()
        view_options = ["Weekly Intelligence Brief", "Trend Timelines", "Archives", "Clean Archives", "Trash"]
    else:
        st.sidebar.info("Access Level: **Guest Viewer**")
        st.sidebar.caption("*(System Access Restricted)*")
        view_options = ["Weekly Intelligence Brief", "Trend Timelines", "Archives"]

    st.sidebar.markdown("---")
    view_selection = st.sidebar.radio("Select View:", view_options)

    if advanced_tool == "Quantitative Threat Scoring":
        st.title("Quantitative Threat Scoring")
        st.markdown("Algorithmic ranking of regional supply chain vulnerability based on historical archive data.")
        
        archive_mapping = get_brief_mappings('data')
        if archive_mapping:
            scoring_data = {}
            for f_path in archive_mapping.values():
                try:
                    with open(f_path, 'r') as file:
                        d = json.load(file)
                        actions = d.get('recent_actions', [])
                        
                        for action in actions:
                            loc = str(action.get('Location', '')).strip()
                            if loc and loc != "Global":
                                for country, data_tuple in COUNTRY_INFO.items():
                                    if country.lower() in loc.lower():
                                        region = data_tuple[1]
                                        scoring_data[region] = scoring_data.get(region, 0) + 1
                                        break
                except: pass
            
            if scoring_data:
                score_df = pd.DataFrame(list(scoring_data.items()), columns=["Region", "Instability Actions Logged"])
                score_df = score_df.sort_values(by="Instability Actions Logged", ascending=False).reset_index(drop=True)
                
                def assign_threat(score):
                    if score > 10: return "🔴 Critical"
                    elif score > 5: return "🟠 High"
                    elif score > 2: return "🟡 Elevated"
                    else: return "🟢 Standard"
                    
                score_df["Calculated Threat Level"] = score_df["Instability Actions Logged"].apply(assign_threat)
                
                st.table(score_df.set_index(score_df.columns[0]))
            else:
                st.warning("Not enough historical data to generate scores yet.")
        else:
            st.warning("No archives available.")
            
        st.stop() 
        
    elif advanced_tool == "Intelligence Interrogation (RAG)":
        st.title("Intelligence Interrogation (RAG)")
        st.markdown("Query the historical SemicoN database. Responses are generated strictly from your vetted archives.")

        if not client:
            st.error("⚠️ GEMINI_API_KEY is missing from Koyeb Environment Variables. Please add it to unlock this feature.")
            st.stop()

        if "rag_messages" not in st.session_state:
            st.session_state.rag_messages = []

        for message in st.session_state.rag_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if prompt := st.chat_input("Ask a strategic question (e.g., 'What were the major export controls on China last month?'):"):
            st.session_state.rag_messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                message_placeholder.markdown("Scanning intelligence archives & calculating threat models... 🕵️‍♂️")

                archive_mapping = get_brief_mappings('data')
                context_data = ""
                
                # --- RAG 2.0 CONTEXT BUILDER (OPTIMIZED FOR SPEED & ACCURACY) ---
                user_keywords = [w.lower() for w in re.findall(r'\b\w+\b', prompt) if len(w) > 2 and w.lower() not in ['what', 'when', 'where', 'which', 'who', 'why', 'how', 'were', 'was', 'this', 'that', 'with', 'from', 'about', 'the', 'and', 'for', 'are', 'did', 'have', 'has']]

                file_scores = []
                for f_path in archive_mapping.values():
                    try:
                        with open(f_path, 'r') as file:
                            d = json.load(file)
                            content = d.get('brief_raw', '').lower() + json.dumps(d.get('recent_actions', [])).lower()
                            score = sum(content.count(kw) for kw in user_keywords)
                            file_scores.append((score, f_path))
                    except: pass

                # Sort files by relevance score
                file_scores.sort(key=lambda x: x[0], reverse=True)

                # Pick top 3 most relevant files, fallback to latest 2 if no keywords match
                top_files = [fs[1] for fs in file_scores if fs[0] > 0][:3]
                if not top_files:
                    top_files = list(archive_mapping.values())[:2]
                
                for f_path in top_files:
                    try:
                        with open(f_path, 'r') as file:
                            d = json.load(file)
                            r_text = d.get('brief_raw', '')
                            
                            categories = [
                                ("Global Foundry Market", extract_tag('EXEC', r_text) or ""),
                                ("AI Chip Demand", extract_tag('LITHO', r_text) or ""),
                                ("Critical Minerals (REE)", extract_tag('REE', r_text) or ""),
                                ("Export Controls", extract_tag('GEO', r_text) or ""),
                                ("Military & Outer Space", extract_tag('MILITARY', r_text) or ""),
                                ("India Developments", extract_tag('INDIA', r_text) or ""),
                                ("West Asia / Middle East", extract_tag('WEST_ASIA', r_text) or "")
                            ]
                            
                            context_data += f"\n\n--- INTELLIGENCE BRIEF DATE: {d.get('date', 'Unknown')} ---\n"
                            context_data += "ALGORITHMIC THREAT SCORES:\n"
                            
                            for name, txt in categories:
                                if len(txt.strip()) > 20:
                                    score = calculate_domain_threat(name, txt, d)
                                    context_data += f"- {name}: {score}%\n"
                                    
                            context_data += "\nRAW INTELLIGENCE TEXT:\n"
                            context_data += r_text
                            context_data += f"\nLOGGED STATE ACTIONS:\n{json.dumps(d.get('recent_actions', []))}"
                    except: pass

                # --- RAG 2.0 SYSTEM PROMPT ---
                sys_prompt = f"""
                You are an elite geopolitical intelligence AI assistant for the SemicoN Dashboard.
                Your primary directive is to answer the user's question using ONLY the provided historical intelligence archives below.
                
                CRITICAL RAG 2.0 DIRECTIVE: You now have access to the "Algorithmic Threat Scores" (0-100%) calculated for each domain.
                - Scores > 60% indicate High Volatility/Threat environments.
                - Scores < 40% indicate standard baseline risk.
                When analyzing threats, vulnerabilities, or risks, explicitly cite these Algorithmic Threat Scores to ground your reasoning in the quantitative data model.
                
                If the answer is not present in the context, explicitly state: "I cannot find this information in the vetted intelligence archives."
                Format your responses in a structured, highly analytical style. Use structured comparative tables if you are summarizing multiple actors, dates, or data points. Do NOT invent or hallucinate external information.
                
                ARCHIVES CONTEXT:
                {context_data}
                """
                
                try:
                    full_response = ""
                    response = client.models.generate_content_stream(
                        model=model_name,
                        contents=[sys_prompt, prompt]
                    )
                    for chunk in response:
                        if chunk.text:
                            full_response += chunk.text
                            message_placeholder.markdown(full_response + "▌") 
                            
                    message_placeholder.markdown(full_response)
                except Exception as e:
                    full_response = f"⚠️ Error querying the intelligence database: {e}"
                    message_placeholder.markdown(full_response)
            st.session_state.rag_messages.append({"role": "assistant", "content": full_response})
            
        st.stop() 

    if view_selection == "Weekly Intelligence Brief":
        is_editing = st.session_state.get('vetting_toggle', False)
            
        if is_editing and st.session_state['role'] == 'admin':
            st.warning("You are currently in Edit Mode. Changes made here will be permanently written to the intelligence database.")
            with st.form("editor_form"):
                st.markdown("### Edit Intelligence Text")
                new_sum = st.text_area("Executive Summary", text_summary, height=150)
                new_ews = st.text_area("Early Warning & Red Flags", text_ews, height=100)
                new_s1 = st.text_area("Global Foundry Market", text_section_1, height=150)
                new_s2 = st.text_area("AI Chip Demand", text_section_2, height=100)
                new_s3 = st.text_area("Critical Minerals", text_section_3, height=100)
                new_s4 = st.text_area("Export Controls", text_section_4, height=100)
                new_mil = st.text_area("Military & Outer Space", text_military, height=100)
                new_s5 = st.text_area("Lithography Chokepoints", text_section_5, height=150)
                new_india = st.text_area("India: Domestic & Strategic Developments", text_india, height=150)
                new_wa = st.text_area("West Asia: Domestic & Strategic Developments", text_wa, height=150)
                new_fin = st.text_area("Strategic Conclusion", text_final, height=150)
                
                st.markdown("### Edit Data Tables & KPIs")
                
                edited_kpi_df = st.data_editor(clean_dataframe(pd.DataFrame(dashboard_data.get('kpi_metrics', []))), num_rows="dynamic", use_container_width=True)
                edited_funding_df = st.data_editor(clean_dataframe(pd.DataFrame(dashboard_data.get('funding_data', []))), num_rows="dynamic", use_container_width=True)
                edited_market_df = st.data_editor(clean_dataframe(pd.DataFrame(dashboard_data.get('market_impact', []))), num_rows="dynamic", use_container_width=True)
                edited_risk_df = st.data_editor(clean_dataframe(pd.DataFrame(dashboard_data.get('supply_chain_risk', []))), num_rows="dynamic", use_container_width=True)
                edited_actions_df = st.data_editor(clean_dataframe(pd.DataFrame(dashboard_data.get('recent_actions', []))), num_rows="dynamic", use_container_width=True)
                
                if st.form_submit_button("💾 Save Vetted Intelligence"):
                    
                    clean_kpi = edited_kpi_df.fillna("")
                    clean_actions = edited_actions_df.fillna("")
                    clean_fund = edited_funding_df.fillna("")
                    clean_market = edited_market_df.fillna("")
                    clean_risk = edited_risk_df.fillna("")
                    
                    kpi_json = clean_kpi.to_json(orient='records')
                    matrix_json = clean_actions.to_json(orient='records')
                    fund_json = clean_fund.to_json(orient='records')
                    market_json = clean_market.to_json(orient='records')
                    risk_json = clean_risk.to_json(orient='records')
                    
                    new_raw = f"<SUMMARY>\n{new_sum}\n</SUMMARY>\n\n<EWS>\n{new_ews}\n</EWS>\n\n<EXEC>\n{new_s1}\n</EXEC>\n\n<LITHO>\n{new_s2}\n</LITHO>\n\n<REE>\n{new_s3}\n</REE>\n\n<GEO>\n{new_s4}\n</GEO>\n\n<MILITARY>\n{new_mil}\n</MILITARY>\n\n<CONCLUSION>\n{new_s5}\n</CONCLUSION>\n\n<INDIA>\n{new_india}\n</INDIA>\n\n<WEST_ASIA>\n{new_wa}\n</WEST_ASIA>\n\n<FINAL_CONCLUSION>\n{new_fin}\n</FINAL_CONCLUSION>\n\n<KPI_METRICS>{kpi_json}</KPI_METRICS>\n<FUNDING_DATA>{fund_json}</FUNDING_DATA>\n<MARKET_IMPACT>{market_json}</MARKET_IMPACT>\n<RISK_INDEX>{risk_json}</RISK_INDEX>\n<ACTION_MATRIX>{matrix_json}</ACTION_MATRIX>"
                    
                    dashboard_data['brief_raw'] = new_raw
                    dashboard_data['kpi_metrics'] = clean_kpi.to_dict(orient='records')
                    dashboard_data['recent_actions'] = clean_actions.to_dict(orient='records')
                    dashboard_data['funding_data'] = clean_fund.to_dict(orient='records')
                    dashboard_data['market_impact'] = clean_market.to_dict(orient='records')
                    dashboard_data['supply_chain_risk'] = clean_risk.to_dict(orient='records')
                    
                    if latest_filepath:
                        with open(latest_filepath, 'w') as f:
                            json.dump(dashboard_data, f)
                        
                    st.cache_data.clear() 
                    st.success("Changes permanently saved! Refreshing system...")
                    time.sleep(1)
                    st.rerun() 
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.toggle("✏️ Enable Human-AI Vetting Mode (Edit Text & Data)", key="vetting_toggle")

        else:
            st.title(f"SemicoN Weekly Brief - {brief_date}") 
            st.markdown("---")
            
            current_day = datetime.now(timezone.utc).astimezone().strftime('%B %d, %Y')

            # Change margin-bottom: 10px; to margin-bottom: 30px;
            st.markdown(f"<h3 style='color:#ff4b4b; margin-top: 10px; margin-bottom: 30px;'>🛡️ Strategic Threat Monitor ({current_day})</h3>", unsafe_allow_html=True)
            
            check_early_warnings()

            st.markdown("<br>", unsafe_allow_html=True)
            run_shockwave_engine()
            st.markdown("<br><hr><br>", unsafe_allow_html=True)

        # --- NEW: LIVE GLOBAL SEMICONDUCTOR RISK INDEX (24-HOUR SYNC) ---
            # 1. Get baseline structural risk from the weekly SCV domains
            all_texts = [text_section_1, text_section_2, text_section_3, text_section_4, text_military, text_india, text_wa]
            valid_scores = [calculate_domain_threat("domain", t, dashboard_data) for t in all_texts if len(t.strip()) > 20]
            baseline_risk = int(sum(valid_scores) / len(valid_scores)) if valid_scores else 40

            # 2. Divide collected 24-hour news into thematic categories
            live_rss_data = parse_rss_txt_file()
            theme_scores = {"Kinetic": 0, "Economic": 0, "Supply": 0}
            breaking_news = []
            
            if live_rss_data:
                kw_kinetic = ['war', 'military', 'strike', 'blockade', 'escalation', 'breach', 'crisis']
                kw_economic = ['sanction', 'tariff', 'export control', 'ban', 'subsidy', 'embargo']
                kw_supply = ['shortage', 'disrupt', 'delay', 'chokepoint', 'vulnerability']
                
                for region, articles in live_rss_data.items():
                    for art in articles:
                        if not art.get('is_24h', False): continue # STRICT 24-HOUR FILTER
                        
                        title_lower = art['title'].lower()
                        
                        # --- FIX: ADDED MISSING DEFINITIONS FOR ALL HIT COUNTERS ---
                        k_hits = sum(1 for kw in kw_kinetic if kw in title_lower)
                        e_hits = sum(1 for kw in kw_economic if kw in title_lower)
                        s_hits = sum(1 for kw in kw_supply if kw in title_lower)
                        
                        theme_scores["Kinetic"] += k_hits * 5.0
                        theme_scores["Economic"] += e_hits * 3.5
                        theme_scores["Supply"] += s_hits * 4.0
                        
                        # Capture breaking alert dynamically
                        if (k_hits + e_hits + s_hits) >= 2 and not breaking_news:
                            breaking_news.append(art)

            # 3. Mathematical Method: Weighted Composite Risk Model with Logarithmic Normalization
            import math
            def log_scale(score, max_boost):
                # Asymptotic curve: Prevents linear stacking from maxing out at 100%
                return max_boost * (1 - math.exp(-0.06 * score)) if score > 0 else 0

            # Distribute the maximum allowed volatility (60%) across themes
            kinetic_volatility = log_scale(theme_scores["Kinetic"], 25)  # Hardest impact
            economic_volatility = log_scale(theme_scores["Economic"], 20)
            supply_volatility = log_scale(theme_scores["Supply"], 15)

            # Final Composite Math: 40% Structural Baseline + 60% Dynamic Volatility
            raw_composite = (baseline_risk * 0.4) + kinetic_volatility + economic_volatility + supply_volatility
            global_risk = int(round(raw_composite + 25)) # Apply a standard operational floor
            global_risk = max(20, min(99, global_risk)) # Cap at 99% to maintain metric authenticity

            risk_cols = st.columns([1, 1])
            with risk_cols[0]:
                if global_risk >= 75:
                    st.error(f"🔴 **Global Semiconductor Risk Index – In the Past 24-Hours (Logarithmic Composite Model): {global_risk} / 100** (Critical)")
                elif global_risk >= 50:
                    st.warning(f"🟠 **Global Semiconductor Risk Index – In the Past 24-Hours (Logarithmic Composite Model): {global_risk} / 100** (Rising Risk)")
                else:
                    st.success(f"🟢 **Global Semiconductor Risk Index – In the Past 24-Hours (Logarithmic Composite Model): {global_risk} / 100** (Stable)")

            with risk_cols[1]:
                if breaking_news:
                    st.error(f"🚨 **BREAKING ALERT – In the Past 24-Hours:** [{breaking_news[0]['title']}]({breaking_news[0]['link']})")
                else:
                    st.info("📡 **Live Radar:** No immediate kinetic or economic breaks detected in the past 24-hours.")

            # --- PLOTLY CONCENTRIC RING WHEEL IMPLEMENTATION ---
            st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 30px; margin-bottom: 10px;'>Semicon, Rare Earth and AI Geopolitical Outlook</h3>", unsafe_allow_html=True)

            scv_categories = [
                ("Global Foundry Market", text_section_1, "#00bfff"),
                ("AI Chip Demand", text_section_2, "#ff00ff"),
                ("Critical Minerals (REE)", text_section_3, "#00ff00"),
                ("Export Controls", text_section_4, "#ff4b4b"),
                ("Military & Outer Space", text_military, "#ffd166"),
                ("India Developments", text_india, "#ff8c00"),
                ("West Asia / Middle East", text_wa, "#9400d3")
            ]

            active_cats = []
            for name, txt, col in scv_categories:
                if len(txt.strip()) > 20:
                    dynamic_score = calculate_domain_threat(name, txt, dashboard_data)
                    active_cats.append({"name": name, "score": dynamic_score, "color": col})

            active_cats = sorted(active_cats, key=lambda x: x["score"], reverse=True)

            scv_cols = st.columns([1.2, 1])

            with scv_cols[0]:
                st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
                st.markdown("<p style='color: #888; font-size: 14px; font-weight: bold; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 0px; text-align: center;'>Supply Chain Vulnerability (SCV) Wheel (weekly)</p>", unsafe_allow_html=True)

                if active_cats:
                    fig = go.Figure()

                    base_hole = 0.35      
                    ring_width = 0.015    
                    gap = 0.075          

                    for i, cat in enumerate(active_cats):
                        val = cat["score"]
                        color = cat["color"]

                        data_hole = base_hole + i * (ring_width + gap)
                        fig.add_trace(go.Pie(
                            values=[val, 100 - val],
                            hole=data_hole,
                            domain=dict(x=[0, 1], y=[0, 1]),
                            marker=dict(
                                colors=[color, "#000000"],
                                line=dict(width=0) 
                            ),
                            textinfo='none',
                            sort=False,
                            direction='clockwise',
                            hoverinfo='text',
                            hovertext=[f"{cat['name']}: {val}%", ""],
                            showlegend=False
                        ))

                        gap_hole = data_hole + ring_width
                        fig.add_trace(go.Pie(
                            values=[100], 
                            hole=gap_hole,
                            domain=dict(x=[0, 1], y=[0, 1]),
                            marker=dict(
                                colors=["#000000"], 
                                line=dict(width=2, color="#000000")
                            ), 
                            textinfo='none',
                            sort=False,
                            hoverinfo='none',
                            showlegend=False
                        ))
                    
                    fig.add_trace(go.Pie(
                        values=[100],
                        hole=0.98,
                        domain=dict(x=[0, 1], y=[0, 1]),
                        marker=dict(colors=["#000000"], line=dict(width=4, color="#000000")),
                        textinfo='none',
                        hoverinfo='none',
                        showlegend=False
                    ))

                    overall_score = int(sum(c["score"] for c in active_cats) / len(active_cats)) if active_cats else 0

                    fig.update_layout(
                        paper_bgcolor='rgba(0,0,0,0)',
                        plot_bgcolor='rgba(0,0,0,0)',
                        margin=dict(t=20, b=20, l=20, r=20),
                        height=400,
                        annotations=[dict(
                            text=f"<b style='font-size:42px; color:white;'>{overall_score}</b><br><span style='color:#aaaaaa; font-size:12px; font-weight:bold;'>AVG SCV SCORE</span>",
                            x=0.5, y=0.5,
                            showarrow=False
                        )]
                    )

                    st.plotly_chart(fig, use_container_width=True)
                    
                    legend_html = "<div style='display:flex; flex-wrap:wrap; justify-content:center; gap:12px; margin-top:-30px; margin-bottom:20px;'>"
                    for cat in active_cats:
                        legend_html += f"<div style='font-size:10px; font-weight:bold; color:#a3a3a3;'><span style='color:{cat['color']};'>●</span> {cat['name'].upper()}</div>"
                    legend_html += "</div>"
                    st.markdown(legend_html, unsafe_allow_html=True)
                else:
                    st.warning("Not enough data to render the wheel.")

            with scv_cols[1]:
                st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
                st.markdown("<p style='color: #888; font-size: 14px; font-weight: bold; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 20px;'>SCV Threat Matrix (Active Domains) (Weekly)</p>", unsafe_allow_html=True)
                
                if active_cats:
                    for cat in active_cats:
                        label = cat["name"].upper()
                        value = cat["score"]
                        color = cat["color"]
                        
                        st.markdown(f"""
                        <div style="margin-bottom: 15px;">
                            <div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
                                <span style="font-size: 12px; font-weight: 600; color: #d1d5db; letter-spacing: 0.5px;">{label}</span>
                                <span style="font-size: 13px; font-weight: bold; color: {color};">{value}%</span>
                            </div>
                            <div style="width: 100%; background-color: #1f2937; border-radius: 4px; height: 6px;">
                                <div style="width: {value}%; background-color: {color}; height: 6px; border-radius: 4px; box-shadow: 0 0 8px {color}80;"></div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("No domain data available to calculate threat matrix.")
            
            # ==========================================
            # AI GEOPOLITICAL SYNTHESIS
            # ==========================================
            st.markdown("<p style='color: #888; font-size: 14px; font-weight: bold; text-transform: uppercase; letter-spacing: 1px; margin-top: 20px; margin-bottom: 10px;'>AI Geopolitical Synthesis (Weekly) - The Big Picture</p>", unsafe_allow_html=True)
            if text_summary and text_summary.strip() != "":
                with st.container(height=160):
                    render_highlighted_text(text_summary, selected_actor)
            else:
                st.info("No synthesis data available.")
                
            st.markdown("---")

            if text_ews and text_ews.strip() != "":
                st.markdown("<h5 style='color:#ff4b4b; margin-top: 0px; margin-bottom: 5px;'>🚨 Weekly EWS Synthesis - The Tactical Alarm</h5>", unsafe_allow_html=True)
                render_highlighted_text(text_ews, selected_actor)
                st.markdown("<br>", unsafe_allow_html=True)
            
            st.markdown("---")

            if latest_filepath:
                df_kpi = clean_dataframe(pd.DataFrame(dashboard_data.get('kpi_metrics', [])))
                df_fund = clean_dataframe(pd.DataFrame(dashboard_data.get('funding_data', [])))
                df_market = clean_dataframe(pd.DataFrame(dashboard_data.get('market_impact', [])))
                df_risk = clean_dataframe(pd.DataFrame(dashboard_data.get('supply_chain_risk', [])))
                
                if not df_kpi.empty and 'Metric' in df_kpi.columns:
                    st.markdown("##### Executive Snapshot")
                    cols_per_row = 3
                    for i in range(0, len(df_kpi), cols_per_row):
                        kpi_cols = st.columns(cols_per_row)
                        for j in range(cols_per_row):
                            if i + j < len(df_kpi):
                                with kpi_cols[j]:
                                    row = df_kpi.iloc[i + j]
                                    label_val = str(row.get('Metric', ''))
                                    metric_val = str(row.get('Value', ''))
                                    st.markdown(f"""
                                        <div style="margin-bottom: 20px; border-left: 4px solid #00bfff; padding-left: 15px;">
                                            <p style="font-size: 13px; font-weight: 600; color: #888; margin-bottom: 0px; line-height: 1.2;">{label_val}</p>
                                            <p style="font-size: 26px; font-weight: bold; margin-top: 0px; line-height: 1.2; color: white;">{metric_val}</p>
                                        </div>
                                    """, unsafe_allow_html=True)
                    st.markdown("---")
                
                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Executive Summary</h3>", unsafe_allow_html=True)
                if text_summary: render_highlighted_text(text_summary, selected_actor)

                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Global Foundry Market & Geopolitical Positioning</h3>", unsafe_allow_html=True)
                if text_section_1: render_highlighted_text(text_section_1, selected_actor)
                
                if not df_fund.empty and not (len(df_fund) == 1 and "No " in str(df_fund.iloc[0].values[0])):
                    st.markdown("##### Strategic Investments & Funding")
                    st.table(df_fund.set_index(df_fund.columns[0]))

                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>AI Chip Demand, Manufacturing & Processing</h3>", unsafe_allow_html=True)
                if text_section_2: render_highlighted_text(text_section_2, selected_actor)
                
                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Critical Minerals: Rare Earth Reserves & Supply Chains</h3>", unsafe_allow_html=True)
                if text_section_3: render_highlighted_text(text_section_3, selected_actor)
                
                if not df_market.empty and not (len(df_market) == 1 and "Data Unavailable" in str(df_market.iloc[0].values[0])):
                    st.markdown("##### Market & Geopolitical Impact")
                    st.table(df_market.set_index(df_market.columns[0]))

                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Export Controls & Geopolitical Impact</h3>", unsafe_allow_html=True)
                if text_section_4: render_highlighted_text(text_section_4, selected_actor)
                
                if not df_risk.empty and not (len(df_risk) == 1 and "Data Unavailable" in str(df_risk.iloc[0].values[0])):
                    st.markdown("##### Supply Chain Risk Analysis")
                    st.table(df_risk.set_index(df_risk.columns[0]))
                    
                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>AI, Chips and Rare Earth in Military and Outer Space Domain</h3>", unsafe_allow_html=True)
                if text_military: render_highlighted_text(text_military, selected_actor)
                
                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Lithography Chokepoints & State Actions</h3>", unsafe_allow_html=True)
                if text_section_5: render_highlighted_text(text_section_5, selected_actor)

                if text_india and text_india.strip() != "": 
                    st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>India: Domestic & Strategic Developments</h3>", unsafe_allow_html=True)
                    render_highlighted_text(text_india, selected_actor)
                    
                if text_wa and text_wa.strip() != "": 
                    st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>West Asia/Middle East: Domestic & Strategic Developments</h3>", unsafe_allow_html=True)
                    render_highlighted_text(text_wa, selected_actor)

                region_colors = {
                    "Asia": "#00bfff",                
                    "West Asia/Middle East": "#00ff00",
                    "Americas": "#ff4b4b",            
                    "Africa": "#ffd166",                
                    "Europe": "#ff69b4",                
                    "Oceania": "#ffa500"                
                }

                if not df_actions.empty:
                    map_data = []
                    for _, row in df_actions.iterrows():
                        if selected_actor != "All" and selected_actor != str(row.get('Actor', '')):
                            continue
                            
                        loc_val = str(row.get('Location', '')).strip()
                        actor_val = str(row.get('Actor', ''))
                        action_val = str(row.get('Action', '')).strip()
                        
                        search_string = f"{loc_val} {actor_val}".lower()

                        added_iso = set()
                        
                        for country, data_tuple in COUNTRY_INFO.items():
                            iso_code = data_tuple[0]
                            region_name = data_tuple[1]
                            c_low = country.lower()
                            
                            if re.search(rf'(?<![a-z]){re.escape(c_low)}(?![a-z])', search_string):
                                if iso_code not in added_iso:
                                    map_data.append({"iso_alpha": iso_code, "country": country, "actor": actor_val, "action": action_val, "region": region_name})
                                    added_iso.add(iso_code)
                    
                    if map_data:
                        st.markdown("##### Geopolitical Threat Actions (2D Heatmap)")
                        df_map = pd.DataFrame(map_data)
                        px.set_mapbox_access_token(MAPBOX_PUBLIC_TOKEN)
                        
                        fig = px.choropleth_mapbox(
                            df_map, 
                            geojson="https://raw.githubusercontent.com/johan/world.geo.json/master/countries.geo.json", 
                            locations="iso_alpha", 
                            featureidkey="id", 
                            color="region", 
                            color_discrete_map=region_colors,
                            hover_name="country", 
                            hover_data={"action": True, "actor": True, "region": False, "iso_alpha": False}, 
                            mapbox_style="carto-darkmatter", 
                            zoom=1.0, 
                            center={"lat": 20.0, "lon": 0.0}, 
                            opacity=0.6 
                        )
                        fig.update_layout(margin={"r":0,"t":0,"l":0,"b":0}, paper_bgcolor="rgba(0,0,0,0)", showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)              


#              ==========================================
                    # NEW FEATURE: 2D LIVE SUPPLY CHAIN MAP
                    # ==========================================
                    st.markdown("##### 🌍 Global Semiconductor Infrastructure Map")
                    infra_list = []
                    for category, locations in INFRASTRUCTURE_DATA.items():
                        for loc in locations:
                            infra_list.append({"Facility": loc["name"], "Lat": loc["lat"], "Lon": loc["lon"], "Category": category})
                    
                    if infra_list:
                        df_infra = pd.DataFrame(infra_list)
                        px.set_mapbox_access_token(MAPBOX_PUBLIC_TOKEN)
                        fig_infra = px.scatter_mapbox(
                            df_infra, lat="Lat", lon="Lon", color="Category", hover_name="Facility",
                            mapbox_style="carto-darkmatter", zoom=1.2, height=500,
                            color_discrete_sequence=["#00ffff", "#ff00ff", "#ffff00", "#39ff14", "#ff4500", "#ffffff"]
                        )
                        fig_infra.update_layout(margin={"r":0,"t":0,"l":0,"b":0}, paper_bgcolor="rgba(0,0,0,0)", legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01, font=dict(color="white")))
                        st.plotly_chart(fig_infra, use_container_width=True)
                    st.markdown("<br>", unsafe_allow_html=True)

                    # ==========================================
                    # NEW FEATURE: TRUE 3D MAPBOX GLOBE
                    # ==========================================
                    st.markdown("##### 3D Tactical Infrastructure Globe")
                    st.markdown("<p style='font-size: 13px; color: #888;'>Use <b>Right-Click + Drag</b> to rotate the 3D globe. Scroll to zoom.</p>", unsafe_allow_html=True)
                    
                    selected_infra = st.multiselect(
                        "📍 Toggle Physical Infrastructure Layers:", 
                        [
                            "Semiconductor Fabs", 
                            "Critical Mineral Sites", 
                            "Maritime Chokepoints", 
                            "Gulf FDI & Capital Diplomacy", 
                            "Naval Order of Battle & Strategic Bases",
                            "Aerospace & Space Force Installations" 
                        ], 
                        default=["Semiconductor Fabs", "Maritime Chokepoints"]
                    )
                    
                    infra_colors_hex = {
                        "Semiconductor Fabs": "#00ffff", 
                        "Critical Mineral Sites": "#ff00ff", 
                        "Maritime Chokepoints": "#ffff00",
                        "Gulf FDI & Capital Diplomacy": "#39ff14",
                        "Naval Order of Battle & Strategic Bases": "#ff4500", 
                        "Aerospace & Space Force Installations": "#ffffff" 
                    }

                    # Prepare data to send to JavaScript
                    map_features = []
                    for infra_type in selected_infra:
                        sites = INFRASTRUCTURE_DATA.get(infra_type, [])
                        color = infra_colors_hex[infra_type]
                        for site in sites:
                            map_features.append({
                                "name": site["name"],
                                "lat": site["lat"],
                                "lon": site["lon"],
                                "color": color
                            })

                    map_data_json = json.dumps(map_features)
                    mapbox_token = MAPBOX_PUBLIC_TOKEN
                    map_bg_color = "#000000"
                    
                    html_code = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                    <meta charset="utf-8" />
                    <script src="https://api.mapbox.com/mapbox-gl-js/v2.15.0/mapbox-gl.js"></script>
                    <link href="https://api.mapbox.com/mapbox-gl-js/v2.15.0/mapbox-gl.css" rel="stylesheet" />
                    <style>
                    html, body {{ height: 100%; width: 100%; margin: 0; padding: 0; background-color: {map_bg_color}; overflow: hidden; }}
                    #map {{ position: absolute; top: 0; left: 0; width: 100%; height: 100%; border-radius: 8px; }}
                    .mapboxgl-popup-content {{ background-color: #1e1e1e; color: #ffffff; border: 1px solid #00bfff; border-radius: 5px; font-family: monospace; padding: 10px; box-shadow: 0 0 10px rgba(0, 191, 255, 0.5); }}
                    .mapboxgl-popup-close-button {{ color: #ffffff; }}
                    </style>
                    </head>
                    <body>
                    <div id="map"></div>
                    <script>
                    mapboxgl.accessToken = '{mapbox_token}';

                    if (!mapboxgl.accessToken.startsWith('pk.')) {{
                        document.getElementById('map').innerHTML = '<div style="display:flex; justify-content:center; align-items:center; height:100%; color:#00bfff; font-family:monospace; text-align:center; background:#1e1e1e; border:1px solid #00bfff; border-radius:8px; padding:20px;"><div><h2>⚠️ Mapbox Token Required</h2></div></div>';
                    }} else {{
                        const map = new mapboxgl.Map({{
                            container: 'map',
                            style: 'mapbox://styles/mapbox/dark-v11',
                            projection: 'globe', 
                            zoom: 1.2,
                            center: [30, 20],
                            pitch: 45
                        }});

                        map.on('style.load', () => {{
                            map.setFog({{ 'color': 'rgb(10, 20, 30)', 'high-color': 'rgb(0, 0, 0)', 'horizon-blend': 0.1, 'space-color': 'rgb(5, 5, 5)', 'star-intensity': 0.8 }});
                        }});

                        map.addControl(new mapboxgl.NavigationControl());

                        const rawData = {map_data_json};

                        function createPolygon(lon, lat, radiusDegrees = 1.0) {{
                            const pts = [];
                            const sides = 16;
                            for (let i = 0; i < sides; i++) {{
                                const angle = (i / sides) * 2 * Math.PI;
                                const lonOffset = (radiusDegrees / Math.cos(lat * Math.PI / 180)) * Math.cos(angle);
                                const latOffset = radiusDegrees * Math.sin(angle);
                                pts.push([lon + lonOffset, lat + latOffset]);
                            }}
                            pts.push(pts[0]); 
                            return [pts];
                        }}

                        map.on('load', () => {{
                            map.resize();
                            
                            const features = rawData.map(item => ({{
                                type: 'Feature',
                                geometry: {{ type: 'Polygon', coordinates: createPolygon(item.lon, item.lat, 0.8) }},
                                properties: {{ name: item.name, color: item.color, height: 500000 }}
                            }}));

                            map.addSource('infrastructure', {{ type: 'geojson', data: {{ type: 'FeatureCollection', features: features }} }});
                            map.addLayer({{
                                'id': 'infrastructure-pillars',
                                'type': 'fill-extrusion',
                                'source': 'infrastructure',
                                'paint': {{ 'fill-extrusion-color': ['get', 'color'], 'fill-extrusion-height': ['get', 'height'], 'fill-extrusion-base': 0, 'fill-extrusion-opacity': 0.8 }}
                            }});

                            map.on('click', 'infrastructure-pillars', (e) => {{
                                const props = e.features[0].properties;
                                new mapboxgl.Popup().setLngLat(e.lngLat).setHTML('<strong>' + props.name + '</strong>').addTo(map);
                            }});

                            map.on('mouseenter', 'infrastructure-pillars', () => {{ map.getCanvas().style.cursor = 'pointer'; }});
                            map.on('mouseleave', 'infrastructure-pillars', () => {{ map.getCanvas().style.cursor = ''; }});
                        }});
                    }}
                    </script>
                    </body>
                    </html>
                    """

                    components.html(html_code, height=850, scrolling=False) 

                if not df_actions.empty:
                    st.markdown("##### Recent State Actions")
                    if selected_actor != "All":
                        display_actions = df_actions[df_actions['Actor'] == selected_actor]
                    else:
                        display_actions = df_actions
                        
                    st.table(display_actions.set_index(display_actions.columns[0]))
                    
                    # ==========================================
                    # FEATURE 3 & 4: CORRELATION ENGINE & TIMELINE
                    # ==========================================
                    col_tc1, col_tc2 = st.columns(2)
                    
                    with col_tc1:
                        st.markdown("##### 🔗 Event Correlation Engine - Weekly")
                        st.caption("Detects geographic hotspots where multiple distinct entities are operating simultaneously.")
                        try:
                            loc_group = display_actions.groupby('Location')['Actor'].apply(lambda x: list(set(x))).reset_index()
                            # Strip out vague locations so we only get real strategic convergences
                            loc_group = loc_group[~loc_group['Location'].isin(['Global', 'Multiple', 'Various', 'None', '', 'N/A'])]
                            correlated = loc_group[loc_group['Actor'].apply(len) > 1]
                            
                            if not correlated.empty:
                                for _, row in correlated.head(4).iterrows():
                                    actors_str = ", ".join(row['Actor'])
                                    loc = row['Location']
                                    
                                    # Updated 4WH Explanatory Format (Natural Language)
                                    html_box = f'''
                                    <div style="background-color: rgba(255,140,0, 0.1); border-left: 3px solid #ff8c00; padding: 12px; margin-bottom: 12px; border-radius: 4px;">
                                        <p style="margin: 0 0 6px 0; font-size: 14px; color: #ff8c00;"><b>⚠️ Strategic Convergence Detected</b></p>
                                        <p style="margin: 0 0 8px 0; font-size: 13px; color: #d1d5db; line-height: 1.5;">
                                            Currently, multiple distinct entities—specifically <b>{actors_str}</b>—are actively concentrating their operational and policy vectors within <b>{loc}</b>, indicating a high-priority strategic hotspot.
                                        </p>
                                        <p style="margin: 0 0 2px 0; font-size: 12px; color: #aaaaaa;"><b>Location:</b> <span style="color: #ddd;">{loc}</span></p>
                                        <p style="margin: 0; font-size: 12px; color: #aaaaaa;"><b>Actors Involved:</b> <span style="color: #ddd;">{actors_str}</span></p>
                                    </div>
                                    '''
                                    st.markdown(html_box, unsafe_allow_html=True)
                            else:
                                st.info("No localized strategic convergences detected this week.")
                        except Exception:
                            pass
                            
                    with col_tc2:
                        st.markdown("##### ⏳ Strategic Timeline Reconstruction - Weekly")
                        try:
                            if len(display_actions) > 0:
                                timeline_html = '<div style="border-left: 2px solid #333; padding-left: 15px; margin-left: 10px;">'
                                for _, row in display_actions.head(5).iterrows():
                                    actor = row.get('Actor', 'Unknown')
                                    # FIX: Removed the [:100] + "..." cutoff. Now the full sentence will render naturally.
                                    action = str(row.get('Action', ''))
                                    
                                    timeline_html += f'<div style="position: relative; margin-bottom: 15px;"><span style="position: absolute; left: -21px; top: 0px; height: 10px; width: 10px; border-radius: 50%; background-color: #00bfff;"></span><div style="font-size: 12px; font-weight: bold; color: #00bfff;">{actor}</div><div style="font-size: 14px; color: #eee; line-height: 1.3; padding-top: 2px;">{action}</div></div>'
                                timeline_html += '</div>'
                                st.markdown(timeline_html, unsafe_allow_html=True)
                        except Exception:
                            pass


                st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 0px;'>Strategic Conclusion</h3>", unsafe_allow_html=True)
                if text_final: render_highlighted_text(text_final, selected_actor)

                # ==========================================
                # LIVE GLOBAL TELEMETRY 
                # ==========================================
                st.markdown("<h2 style='text-align: center; color: #00bfff; margin-top: 50px; margin-bottom: 10px;'>Live Global Telemetry</h2>", unsafe_allow_html=True)
                st.markdown("---")

                # ==========================================
                # NEW: ADVANCED THREAT ANALYTICS 
                # ==========================================
                st.markdown("<h3 style='color:#ff4b4b; font-size:22px; margin-top: 20px; margin-bottom: 10px;'>Advanced Threat Analytics</h3>", unsafe_allow_html=True)

                adv_cols = st.columns([1, 1])

                with adv_cols[0]:
                    # 1. Semiconductor Supply Chain Disruption Monitor
                    st.markdown("##### 🛰️ Supply Chain Disruption Monitor – In the Past 24-Hours")
                    
                    # Dynamically adjust mock risk based on your existing global_risk variable
                    tsmc_risk = "🔴 Critical" if global_risk > 70 else "🟠 Elevated Risk"
                    asml_risk = "🟠 Elevated Risk" if global_risk > 60 else "🟡 Watch"
                    smic_risk = "🔴 Critical" if "china" in str(breaking_news).lower() else "🟠 Elevated Risk"

                    st.markdown(f"""
                    <div style="background-color: #111; padding: 15px; border-radius: 8px; border-left: 4px solid #00bfff;">
                        <p style="margin: 5px 0; color: #ddd;"><strong>TSMC (Taiwan):</strong> {tsmc_risk}</p>
                        <p style="margin: 5px 0; color: #ddd;"><strong>Samsung (Korea):</strong> 🟡 Watch</p>
                        <p style="margin: 5px 0; color: #ddd;"><strong>ASML (Netherlands):</strong> {asml_risk}</p>
                        <p style="margin: 5px 0; color: #ddd;"><strong>SMIC (China):</strong> {smic_risk}</p>
                        <p style="margin: 5px 0; color: #ddd;"><strong>Intel (US):</strong> 🟢 Stable</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # 2. Threat Trend Forecast (Machine Learning Projection)
                    # INCREASED GAP FOR MOBILE RESPONSIVENESS
                    st.markdown("<div style='margin-top: 40px;'></div>", unsafe_allow_html=True)
                    st.markdown("##### 📉 Threat Trend Forecast – In the Past 24-Hours (Machine Learning Projection)")
                    
                    # True Machine Learning (Linear Regression)
                    import numpy as np
                    
                    # Create a time-series model bridging historical baseline to today's live volatility.
                    x_hist = np.array([-5, -4, -3, -2, -1, 0]) # T-minus days
                    y_hist = np.array([max(0, baseline_risk-10), max(0, baseline_risk-5), max(0, baseline_risk-2), min(100, baseline_risk+2), baseline_risk, global_risk])
                    
                    # Fit a 1st-degree polynomial (Linear Regression Line: y = mx + c)
                    ml_model = np.polyfit(x_hist, y_hist, 1)
                    slope = ml_model[0]
                    intercept = ml_model[1]
                    
                    # Calculate RAW mathematical projections (Uncapped)
                    raw_t_plus_3 = int((slope * 3) + intercept)
                    raw_t_plus_7 = int((slope * 7) + intercept)

                    # Cap values for the standard UI display
                    t_plus_3 = min(100, max(0, raw_t_plus_3))
                    t_plus_7 = min(100, max(0, raw_t_plus_7))

                    # Inject CSS for the blinking effect once at the top of the block
                    st.markdown("""
                    <style>
                    @keyframes blinker {
                      50% { opacity: 0; }
                    }
                    </style>
                    """, unsafe_allow_html=True)

                    f_cols = st.columns(3)
                    
                    # Column 1: Today
                    with f_cols[0]:
                        st.metric("Today", f"{global_risk}/100")
                        
                    # Column 2: +3 Days
                    with f_cols[1]:
                        st.metric("+3 Days", f"{t_plus_3}/100", f"{t_plus_3 - global_risk} pts", delta_color="inverse")
                        if raw_t_plus_3 > 100:
                            # REDUCED BLINKING SPEED TO 2.5s
                            st.markdown(f"""
                            <div style="animation: blinker 2.5s linear infinite; color: #ff8c00; font-size: 13px; font-weight: bold; margin-top: -15px;">
                                ⚠️ WATCH: {raw_t_plus_3}% Trajectory
                            </div>
                            """, unsafe_allow_html=True)

                    # Column 3: +7 Days
                    with f_cols[2]:
                        st.metric("+7 Days", f"{t_plus_7}/100", f"{t_plus_7 - global_risk} pts", delta_color="inverse")
                        if raw_t_plus_7 > 100:
                            # If it's way over 100, upgrade from Orange "WATCH" to Red "CRITICAL ALERT"
                            alert_color = "#ff4b4b" if raw_t_plus_7 > 110 else "#ff8c00"
                            alert_text = "🚨 CRITICAL ALERT:" if raw_t_plus_7 > 110 else "⚠️ WATCH:"
                            
                            # REDUCED BLINKING SPEED TO 2s
                            st.markdown(f"""
                            <div style="animation: blinker 2s linear infinite; color: {alert_color}; font-size: 13px; font-weight: bold; margin-top: -15px;">
                                {alert_text} {raw_t_plus_7}% Trajectory
                            </div>
                            """, unsafe_allow_html=True)

                with adv_cols[1]:
                    # 3. Strategic Scenario Simulator
                    st.markdown("##### 🧠 Strategic Scenario Simulator")
                    scenario = st.selectbox("Select Geopolitical Trigger:",
                        ["Taiwan Strait Naval Blockade",
                         "China Rare Earth Export Ban",
                         "US Revokes ASML Servicing Licenses",
                         "Middle East Logistics Chokepoint (Red Sea)"]
                    )

                    if scenario == "Taiwan Strait Naval Blockade":
                        impact = "Chip Supply: -37%<br>AI Hardware Cost: +22%<br>US-China Tension: Extreme"
                        color = "#ff4b4b"
                    elif scenario == "China Rare Earth Export Ban":
                        impact = "REE Supply: -60%<br>EV/Defense Mfg: Critical Delay<br>Global Tension: High"
                        color = "#ff8c00"
                    elif scenario == "US Revokes ASML Servicing Licenses":
                        impact = "China Legacy Chip Cap: -40%<br>ASML Rev: -15%<br>Tech War Tension: High"
                        color = "#ff00ff"
                    else:
                        impact = "Shipping Costs: +300%<br>Logistics Delay: +14 Days<br>Market Tension: Elevated"
                        color = "#ffd166"

                    st.markdown(f"""
                    <div style="background-color: rgba(255, 255, 255, 0.05); border: 1px solid {color}; padding: 15px; border-radius: 8px; margin-top: 10px;">
                        <h6 style="color: {color}; margin-top: 0; font-size: 14px;">Projected Scenario Effects:</h6>
                        <p style="font-family: monospace; color: #ddd; margin-bottom: 0; font-size: 13px; line-height: 1.6;">{impact}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # 4. Strategic Threat Radar Chart
                    # INCREASED GAP FOR MOBILE RESPONSIVENESS
                    st.markdown("<div style='margin-top: 40px;'></div>", unsafe_allow_html=True)
                    st.markdown("##### 📊 Strategic Threat Radar – In the Past 24-Hours (Heuristic Math Method)")
                    
                    # FIX: Define live_volatility by summing your dynamic thematic scores 
                    # before passing it into the radar array
                    live_volatility = kinetic_volatility + economic_volatility + supply_volatility
                    
                    # Generate dynamic radar data based on your existing baseline
                    radar_data = [
                        min(100, global_risk + 5), 
                        min(100, baseline_risk + 15), 
                        min(100, baseline_risk - 5), 
                        min(100, baseline_risk + 20), 
                        min(100, int(live_volatility * 10) + 30)
                    ]
                    
                    radar_fig = go.Figure()
                    radar_fig.add_trace(go.Scatterpolar(
                        r=radar_data,
                        theta=['Export Controls', 'Military Escalation', 'AI Competition', 'Rare Earth Supply', 'Trade War'],
                        fill='toself',
                        line_color='#00bfff',
                        fillcolor='rgba(0, 191, 255, 0.3)'
                    ))
                    radar_fig.update_layout(
                        polar=dict(
                            radialaxis=dict(visible=True, range=[0, 100], color='#888', gridcolor='#333'),
                            angularaxis=dict(color='white', gridcolor='#333')
                        ),
                        showlegend=False,
                        paper_bgcolor='rgba(0,0,0,0)',
                        plot_bgcolor='rgba(0,0,0,0)',
                        margin=dict(t=20, b=20, l=40, r=40),
                        height=250
                    )
                    st.plotly_chart(radar_fig, use_container_width=True)

            # --- NEW: AI NEWS SUMMARY (TOP 10) ---
                sum_cols = st.columns([1.5, 1])
                
                with sum_cols[0]:
                    st.markdown("##### AI Intelligence Summary (Top Radar Hits) – In the Past 24-Hours")
                    if live_rss_data:
                        all_news = []
                        for reg, arts in live_rss_data.items():
                            for art in arts:
                                if art.get('is_24h', False): # STRICT 24-HOUR FILTER
                                    all_news.append(art)
                                    
                        # Deduplicate and sort by threat keyword presence
                        unique_news = {v['title']:v for v in all_news}.values()
                        critical_kw = ['ban', 'sanction', 'shortage', 'escalation', 'military', 'war']
                        sorted_news = sorted(unique_news, key=lambda x: sum(1 for kw in critical_kw if kw in x['title'].lower()), reverse=True)
                        
                        summary_html = "<div style='background-color: #111; padding: 15px; border-radius: 8px; border-left: 4px solid #00bfff; margin-bottom: 20px;'>"
                        for idx, art in enumerate(list(sorted_news)[:10]):
                            summary_html += f"<p style='margin: 5px 0; font-size: 14px;'><span style='color: #00bfff; font-weight: bold;'>{idx+1}.</span> <a href='{art['link']}' target='_blank' style='color: #ddd; text-decoration: none;'>{art['title']}</a></p>"
                        summary_html += "</div>"
                        st.markdown(summary_html, unsafe_allow_html=True)
                
                with sum_cols[1]:
                    # ==========================================
                    # FEATURE 1: INTELLIGENCE SIGNAL DETECTION
                    # ==========================================
                    st.markdown("""
                    <div style="margin-bottom: 15px;">
                        <div style="font-size: 18px; font-weight: 600; margin-bottom: 2px; color: white;">
                            📡 Intelligence Signal Detection Engine – In the Past 24-Hours
                        </div>
                        <div style="font-size: 14px; color: #888888; line-height: 1.4;">
                            Automated scan of live global news tracking secondary geopolitical keywords to surface early 'weak signals' before they escalate into critical threats.
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if live_rss_data:
                        medium_keywords = ['subsidy', 'invest', 'shift', 'policy', 'regulate', 'pressure', 'delay']
                        keyword_counts = {}
                        
                        # Flatten RSS
                        for reg, arts in live_rss_data.items():
                            for art in arts:
                                if not art.get('is_24h', False): continue # STRICT 24-HOUR FILTER
                                
                                title_lower = art['title'].lower()
                                for kw in medium_keywords:
                                    if kw in title_lower:
                                        keyword_counts[kw] = keyword_counts.get(kw, 0) + 1
                        
                        # Find rising signals that aren't critical yet
                        weak_signals = {k: v for k, v in keyword_counts.items() if v > 1}
                        
                        if weak_signals:
                            for signal_kw, count in sorted(weak_signals.items(), key=lambda x: x[1], reverse=True)[:3]:
                                confidence = min(99, count * 15 + 30)
                                st.markdown(f"""
                                <div style="background-color: rgba(255, 255, 0, 0.05); border: 1px solid #ffeb3b; padding: 12px; border-radius: 8px; margin-bottom: 10px;">
                                    <h6 style="color: #ffeb3b; margin: 0 0 5px 0; font-size: 13px;">⚠ Weak Signal Detected</h6>
                                    <p style="margin: 0; font-size: 13px; color: #ddd;">Increased diplomatic/media chatter regarding: <b>{signal_kw.upper()}</b></p>
                                    <p style="margin: 5px 0 0 0; font-size: 11px; color: #888; font-family: monospace;">Confidence: {confidence}% | Trend: Rising ▲</p>
                                </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.info("No abnormal weak signal clusters detected in the past 24-hours.")

                last_updated_str = "Unknown"
                latest_time = 0
                
                for f in glob.glob("data/*"):
                    try:
                        mtime = os.path.getmtime(f)
                        if mtime > latest_time:
                            latest_time = mtime
                    except: pass
                
                if latest_time > 0:
                    diff_minutes = int(time.time() - latest_time) // 60
                    if diff_minutes < 0 or diff_minutes > 100000:
                        last_updated_str = "Awaiting fresh sync"
                    elif diff_minutes < 1:
                        last_updated_str = "Just now"
                    elif diff_minutes < 60:
                        last_updated_str = f"{diff_minutes} minutes ago"
                    else:
                        hours = diff_minutes // 60
                        mins = diff_minutes % 60
                        last_updated_str = f"{hours}h {mins}m ago"

                st.markdown("##### Think Tank Radar")
                st.markdown(f"<p style='font-size: 14px; font-weight: bold; color: #00ff00; margin-top: -10px;'>🟢 LIVE Geopolitical and SemicoN Dashboard SYNC: <span style='color: #888; font-weight: normal;'>Last updated {last_updated_str}</span></p>", unsafe_allow_html=True)
                
                live_rss = parse_rss_txt_file()

                active_alert = get_active_live_alert()
                if active_alert and isinstance(live_rss, dict):
                    alert_headline = active_alert.get('headline', '')
                    
                    is_duplicate = False
                    for reg, articles in live_rss.items():
                        if any(alert_headline.lower() in art.get('title', '').lower() for art in articles):
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        target_reg = "Asia"
                        summary_text = active_alert.get('summary', '').lower()
                        regions_to_check = ["Asia", "West Asia/Middle East", "Americas", "Africa", "Europe", "Oceania"]
                        
                        for r in regions_to_check:
                            if r.split('/')[0].lower() in summary_text:
                                target_reg = r
                                break
                                
                        if target_reg not in live_rss:
                            live_rss[target_reg] = []
                            
                        import urllib.parse
                        # Search ONLY the headline, ignoring the red emoji and warning text
                        alert_query = urllib.parse.quote_plus(alert_headline)
                        
                        live_rss[target_reg].append({
                            "title": f"🔴 LIVE WARNING: {alert_headline}",
                            "published": "JUST IN - ACTIVE ALERT",
                            "link": f"https://news.google.com/search?q={alert_query}"
                        })
                
                if live_rss and isinstance(live_rss, dict):
                    regions = ["Asia", "West Asia/Middle East", "Americas", "Africa", "Europe", "Oceania"]
                    
                    cols_r1 = st.columns(3)
                    for i in range(3):
                        reg = regions[i]
                        color = region_colors.get(reg, "#ffffff")
                        articles = live_rss.get(reg, [])
                        
                        with cols_r1[i]:
                            st.markdown(f"<h4 style='color: {color}; border-bottom: 2px solid {color}; padding-bottom: 5px;'>{reg}</h4>", unsafe_allow_html=True)
                            if articles:
                                scroll_box = st.container(height=300)
                                # Display up to the 8 most recent to keep UI clean but populated
                                for art in list(reversed(articles))[:8]: 
                                    clean_title = art['title'].replace('"', '&quot;').replace("'", "&#39;")
                                    html_str = f'<div style="margin-bottom:10px; padding:10px; background-color:rgba(255,255,255,0.05); border-left:3px solid {color}; border-radius:4px;"><a href="{art["link"]}" target="_blank" style="color:#e0e0e0; font-weight:600; text-decoration:none; font-size:13px; display:block; margin-bottom:5px;">{clean_title}</a><span style="font-size:11px; color:#888;">{art["published"][:25]}</span></div>'
                                    scroll_box.markdown(html_str, unsafe_allow_html=True)
                            else:
                                st.info("No data available.")
                                
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    cols_r2 = st.columns(3)
                    for i in range(3, 6):
                        reg = regions[i]
                        color = region_colors.get(reg, "#ffffff")
                        articles = live_rss.get(reg, [])
                        
                        with cols_r2[i-3]:
                            st.markdown(f"<h4 style='color: {color}; border-bottom: 2px solid {color}; padding-bottom: 5px;'>{reg}</h4>", unsafe_allow_html=True)
                            if articles:
                                scroll_box = st.container(height=300)
                                for art in reversed(articles):
                                    clean_title = art['title'].replace('"', '&quot;').replace("'", "&#39;")
                                    html_str = f'<div style="margin-bottom:10px; padding:10px; background-color:rgba(255,255,255,0.05); border-left:3px solid {color}; border-radius:4px;"><a href="{art["link"]}" target="_blank" style="color:#e0e0e0; font-weight:600; text-decoration:none; font-size:13px; display:block; margin-bottom:5px;">{clean_title}</a><span style="font-size:11px; color:#888;">{art["published"][:25]}</span></div>'
                                    scroll_box.markdown(html_str, unsafe_allow_html=True)
                            else:
                                st.info("No data available.")
                else:
                    st.info("Live feed currently unavailable. (Awaiting next GitHub Action run to initialize new structure).")

                sources_list = dashboard_data.get('sources', [])
                if sources_list:
                    st.markdown("---")
                    st.markdown("<h3 style='color:#00bfff; font-size:22px; margin-top: 20px; margin-bottom: 15px;'>Verified Intelligence Sources</h3>", unsafe_allow_html=True)
                    
                    # Auto-Categorization Engine for Sources
                    themes = {
                        "Global Foundry Market": {"keywords": ["foundry", "tsmc", "samsung", "intel", "smic", "fab", "manufacturing", "yield", "semiconductor", "chipmaker"], "color": "#00bfff", "icon": "🏭", "sources": []},
                        "AI Chip Demand": {"keywords": ["ai", "nvidia", "gpu", "tpu", "compute", "openai", "data center", "server", "algorithm"], "color": "#ff00ff", "icon": "🧠", "sources": []},
                        "Critical Minerals (REE)": {"keywords": ["rare earth", "mineral", "lithium", "cobalt", "graphite", "gallium", "germanium", "mining", "supply chain"], "color": "#00ff00", "icon": "⛏️", "sources": []},
                        "Export Controls & Geopolitics": {"keywords": ["export", "control", "sanction", "ban", "tariff", "entity list", "bis", "geopolitics", "war", "tension", "blockade", "trade"], "color": "#ff4b4b", "icon": "⚖️", "sources": []},
                        "Military & Outer Space": {"keywords": ["military", "defense", "weapon", "missile", "space", "satellite", "darpa", "dod", "navy", "army", "air force", "pentagon"], "color": "#ffd166", "icon": "🚀", "sources": []},
                        "Regional (India & West Asia)": {"keywords": ["india", "modi", "dholera", "tata", "west asia", "middle east", "uae", "saudi", "israel", "gulf", "cg power"], "color": "#ff8c00", "icon": "🌍", "sources": []},
                        "General Strategic Intelligence": {"keywords": [], "color": "#888888", "icon": "📡", "sources": []} # Fallback
                    }

                    for src in sources_list:
                        title_lower = src.get('title', '').lower()
                        placed = False
                        
                        # Keyword matching
                        for t_name, t_data in themes.items():
                            if t_name == "General Strategic Intelligence":
                                continue
                            if any(kw in title_lower for kw in t_data["keywords"]):
                                t_data["sources"].append(src)
                                placed = True
                                break
                        
                        # Fallback for unmatched sources
                        if not placed:
                            themes["General Strategic Intelligence"]["sources"].append(src)

                    # UI Rendering: 2-Column Grid Layout with Custom CSS Cards
                    src_cols = st.columns(2)
                    col_idx = 0
                    
                    for t_name, t_data in themes.items():
                        if t_data["sources"]:
                            with src_cols[col_idx % 2]:
                                theme_html = f"""
                                <div style="background-color: rgba(255,255,255,0.03); border-left: 4px solid {t_data['color']}; padding: 15px; margin-bottom: 15px; border-radius: 6px; height: 100%;">
                                    <h5 style="color: {t_data['color']}; margin-top: 0; margin-bottom: 10px; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px;">{t_data['icon']} {t_name}</h5>
                                    <ul style="margin: 0; padding-left: 20px; font-size: 13px; color: #d1d5db; line-height: 1.6;">
                                """
                                for src in t_data["sources"]:
                                    clean_title = src['title'].replace('"', '&quot;').replace("'", "&#39;")
                                    theme_html += f"<li style='margin-bottom: 5px;'><a href='{src['url']}' target='_blank' style='color: #e0e0e0; text-decoration: none; transition: 0.3s;' onmouseover=\"this.style.color='{t_data['color']}'\" onmouseout=\"this.style.color='#e0e0e0'\">{clean_title}</a></li>"
                                
                                theme_html += "</ul></div>"
                                st.markdown(theme_html, unsafe_allow_html=True)
                            col_idx += 1

            # --- SHARE & EXPORT ---
            st.markdown("---")
            st.markdown("### Document Controls")
            
            if st.session_state['role'] == 'admin':
                bot_col1, bot_col2 = st.columns(2)
                with bot_col1:
                    st.markdown("**🔗 Share Dashboard Link:**")
                    st.code("https://www.semirare.in/", language="text")
                with bot_col2:
                    st.markdown("**📄 Export Full Document:**")
                    if latest_filepath:
                        text_list = [text_summary, text_section_1, text_section_2, text_section_3, text_section_4, text_military, text_section_5, text_india, text_wa]
                        word_file = create_landscape_word(
                            text_list, text_final, 
                            dashboard_data.get('recent_actions', []), 
                            brief_date, 
                            dashboard_data.get('funding_data', []), 
                            dashboard_data.get('market_impact', []), 
                            dashboard_data.get('supply_chain_risk', []),
                            dashboard_data.get('sources', []),
                            text_ews
                        )
                        
                        st.download_button(
                            label="⬇️ Download Authentic Word Doc", data=word_file,
                            file_name=f"SemicoN Weekly Brief - {brief_date}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                if latest_filepath:
                    st.markdown("---")
                    st.markdown("### System Administration")
                    st.toggle("✏️ Enable Human-AI Vetting Mode (Edit Text & Data)", key="vetting_toggle")
            
            else:
                st.markdown("**🔗 Share Dashboard Link:**")
                st.code("https://www.semirare.in/", language="text")

    elif view_selection == "Trend Timelines":
        st.title("Macro Trends & Timelines")
        st.markdown("Tracking the volume of Geopolitical Actions and Supply Chain Risks across historical briefs.")
        
        archive_mapping = get_brief_mappings('data')
        if archive_mapping:
            sorted_brief_paths = list(archive_mapping.values())
            sorted_brief_paths.reverse() 
            
            trend_data = []
            for f_path in sorted_brief_paths:
                try:
                    with open(f_path, 'r') as file:
                        d = json.load(file)
                        b_date = d.get('date', 'Unknown')
                        
                        actions = d.get('recent_actions', [])
                        risks = d.get('supply_chain_risk', [])
                        
                        actions_count = len([a for a in actions if "System" not in str(a.values())])
                        risks_count = len([r for r in risks if "Data Unavailable" not in str(r.values())])
                        
                        trend_data.append({
                            "Date": b_date,
                            "Geopolitical Actions": actions_count,
                            "Supply Chain Risks": risks_count
                        })
                except: pass
            
            if trend_data:
                df_trends = pd.DataFrame(trend_data).set_index("Date")
                st.line_chart(df_trends)
            else:
                st.warning("Not enough valid data points to plot a trend.")
        else:
            st.warning("No archives available to generate trend timelines.")

    elif view_selection == "Archives":
        st.title("Archives")
        archive_mapping = get_brief_mappings('data')
        
        if archive_mapping:
            selected_brief = st.selectbox("Select Past Brief:", list(archive_mapping.keys()))
            st.info(f"Displaying: {selected_brief}")
            with open(archive_mapping[selected_brief], 'r') as f:
                archived_data = json.load(f)
            
            arch_raw = archived_data.get('brief_raw', '')
            t1 = extract_tag('EXEC', arch_raw) or arch_raw
            st.markdown(t1.replace('$', r'\$'))
        else:
            st.warning("No active archives found.")

    elif view_selection == "Clean Archives" and st.session_state['role'] == 'admin':
        st.title("Clean Archives")
        st.info("Move outdated or incorrect briefs to the Trash. Guests cannot see this tool.")
        archive_mapping = get_brief_mappings('data')
        
        if archive_mapping:
            archive_to_clean = st.selectbox("Select Brief to Clean:", list(archive_mapping.keys()))
            if st.button("🗑️ Move to Trash", type="primary"):
                old_path = archive_mapping[archive_to_clean]
                filename = os.path.basename(old_path)
                new_path = os.path.join('trash', filename)
                os.rename(old_path, new_path)
                st.success(f"Successfully moved to Trash: {archive_to_clean}")
                st.rerun()
        else:
            st.warning("No active archives found to clean.")
            
    elif view_selection == "Trash" and st.session_state['role'] == 'admin':
        st.title("Trash Bin")
        st.warning("Items here are completely hidden from the public dashboard.")
        trash_mapping = get_brief_mappings('trash')
        
        if trash_mapping:
            archive_to_manage = st.selectbox("Select Brief in Trash:", list(trash_mapping.keys()))
            old_path = trash_mapping[archive_to_manage]
            filename = os.path.basename(old_path)
            
            col_rev, col_del = st.columns(2)
            with col_rev:
                if st.button("♻️ Revert to Archives"):
                    new_path = os.path.join('data', filename)
                    os.rename(old_path, new_path)
                    st.success(f"Completely restored: {archive_to_manage}")
                    st.rerun()
            with col_del:
                if st.button("⚠️ Permanently Delete", type="primary"):
                    os.remove(old_path)
                    st.success(f"Permanently destroyed: {archive_to_manage}")
                    st.rerun()
        else:
            st.info("Trash is currently empty.")